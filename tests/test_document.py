"""Word 보고서 시험 (요구사항서 10.5).

**여기서 지키는 것 여섯**

    ① 다섯 장이 모두 나온다 — 그리고 수단이 없으면 3·4장을 빼고 번호를 당긴다
    ② 차트 3종이 그림으로 들어간다
    ③ **표가 Word 표 객체다** — 이미지로 넣으면 제안서에 복사해 쓸 수 없다
    ④ 미산출 항목에 빈칸·0원이 아니라 사유가 들어간다
    ⑤ 검토 범위(검토함/미검토)가 마지막 장에 있다
    ⑥ 한글이 깨지지 않는다
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

import pytest
from docx import Document as ReadDocument
from docx.document import Document as DocumentType

from kwise import money
from kwise.compare import ComparisonResult, SensitivityRange
from kwise.diagnose import Diagnosis
from kwise.io import UsageData
from kwise.measures import (
    MEASURE_CATALOG,
    ContractAdjustment,
    PowerFactorResult,
    SolarPoint,
    TariffSwitchResult,
)
from kwise.quality import QualityReport
from kwise.report import (
    CHAPTER_COMPARISON,
    CHAPTER_DIAGNOSIS,
    CHAPTER_MEASURES,
    CHAPTER_SCOPE,
    CHAPTER_SUMMARY,
    DOCUMENT_TITLE,
    DocumentSections,
    build_document,
    document_bytes,
    document_path,
    export_document,
    measure_entries,
)
from kwise.report.document import TABLE_STYLE
from kwise.tariff import BillingResult, TariffTable

# ===================================================================== 도우미


def _style_name(item: object) -> str:
    style = getattr(item, "style", None)
    return str(getattr(style, "name", "")) if style is not None else ""


def _headings(document: DocumentType, level: str = "Heading 1") -> list[str]:
    return [item.text for item in document.paragraphs if _style_name(item) == level]


def _all_text(document: DocumentType) -> str:
    parts = [item.text for item in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _table_with_header(document: DocumentType, *header: str) -> object:
    for table in document.tables:
        if tuple(cell.text for cell in table.rows[0].cells) == header:
            return table
    raise AssertionError(f"머리글 {header} 인 표가 없습니다.")


@pytest.fixture(scope="module")
def entries(
    sample_switch: TariffSwitchResult,
    sample_bill: BillingResult,
    sample_usage: UsageData,
) -> tuple[object, ...]:
    from kwise.measures import evaluate_contract_adjustment

    contract: ContractAdjustment = evaluate_contract_adjustment(
        sample_usage, sample_bill, contract_kw=5_800.0
    )
    return measure_entries(switch=sample_switch, contract=contract)


@pytest.fixture(scope="module")
def full_sections(
    sample_usage: UsageData,
    sample_bill: BillingResult,
    sample_diagnosis: Diagnosis,
    sample_comparison: ComparisonResult,
    entries: tuple[object, ...],
) -> DocumentSections:
    return DocumentSections(
        usage=sample_usage,
        bill=sample_bill,
        diagnosis=sample_diagnosis,
        comparison=sample_comparison,
        measures=entries,  # type: ignore[arg-type]
        building_name="본사 사옥",
        prepared_on=dt.date(2026, 8, 11),
    )


@pytest.fixture(scope="module")
def full_document(full_sections: DocumentSections) -> DocumentType:
    return build_document(full_sections)


@pytest.fixture(scope="module")
def diagnosis_only(
    sample_usage: UsageData, sample_bill: BillingResult, sample_diagnosis: Diagnosis
) -> DocumentType:
    """**수단을 하나도 켜지 않은 보고서.** 진단만 보고 받아 가는 정상 경로다."""
    return build_document(
        DocumentSections(usage=sample_usage, bill=sample_bill, diagnosis=sample_diagnosis)
    )


# ===================================================================== ① 장 구성


def test_다섯_장이_모두_나온다(full_document: DocumentType) -> None:
    """본장 다섯에 **부록 셋**이 이어 붙는다 (22세션 3절)."""
    assert _headings(full_document) == [
        f"1장 {CHAPTER_SUMMARY}",
        f"2장 {CHAPTER_DIAGNOSIS}",
        f"3장 {CHAPTER_MEASURES}",
        f"4장 {CHAPTER_COMPARISON}",
        f"5장 {CHAPTER_SCOPE}",
        "부록 A 산출 근거 상세",
        "부록 B 적용 기준 데이터",
        "부록 C 알려진 한계와 전제",
    ]


def test_표지에_넷이_있다(full_document: DocumentType) -> None:
    text = _all_text(full_document)
    assert DOCUMENT_TITLE in text
    for label in ("건물명", "분석 기간", "작성일", "적용 요금표 시행일"):
        assert label in text
    assert "본사 사옥" in text
    assert "2026-08-11" in text


def test_수단이_없으면_세_장으로_줄고_번호를_당긴다(diagnosis_only: DocumentType) -> None:
    """**비어 있는 장을 제목만 남기지 않는다** — "검토했는데 결과가 없다" 로 읽힌다."""
    assert _headings(diagnosis_only) == [
        f"1장 {CHAPTER_SUMMARY}",
        f"2장 {CHAPTER_DIAGNOSIS}",
        f"3장 {CHAPTER_SCOPE}",
        "부록 B 적용 기준 데이터",
        "부록 C 알려진 한계와 전제",
    ]


def test_수단이_없어도_보고서가_만들어진다(tmp_path: Path, diagnosis_only: DocumentType) -> None:
    """8세션에 Excel 에서 잡은 것과 같은 종류의 결함이다."""
    path = tmp_path / "empty.docx"
    diagnosis_only.save(str(path))
    assert path.stat().st_size > 0
    assert _headings(ReadDocument(str(path)))


def test_각_절이_결론부터_시작한다(full_document: DocumentType) -> None:
    """장 제목 바로 뒤 문단이 결론이다 — 데이터를 앞에 놓지 않는다."""
    paragraphs = full_document.paragraphs
    for index, item in enumerate(paragraphs):
        if _style_name(item) == "Heading 1" and CHAPTER_SUMMARY in item.text:
            first = paragraphs[index + 1]
            assert first.text
            assert any(run.bold for run in first.runs), "결론 문단은 굵게 쓴다."
            break
    else:  # pragma: no cover
        raise AssertionError("요약 장을 찾지 못했습니다.")


# ===================================================================== ② 차트


def test_차트_3종이_삽입된다(diagnosis_only: DocumentType) -> None:
    """부하 프로파일 · 월별 최대수요 · 상위 구간 시각 분포."""
    assert len(diagnosis_only.inline_shapes) == 3
    text = _all_text(diagnosis_only)
    assert "시간대별 평균 부하 프로파일" in text
    assert "월별 최대수요와 요금적용전력" in text
    assert "시각 분포" in text


def test_조합_차트가_더해진다(full_document: DocumentType) -> None:
    # 진단 3장 + 조합 1장 + 수단별 차트 (15세션 2절 — 화면과 같은 프레임을 쓴다).
    assert len(full_document.inline_shapes) >= 4
    assert "조합별 절감액과 투자비" in _all_text(full_document)


def test_수단별_차트가_보고서에도_실린다(full_document: DocumentType) -> None:
    """**새 차트를 Word 에도 반영한다** (15세션 2절 공통 원칙)."""
    text = _all_text(full_document)
    assert "요금제별 요금 구성과 현행 대비 차액" in text


def test_차트가_png_이고_한글_폰트를_쓴다() -> None:
    import matplotlib.pyplot as plt

    from kwise.report.figures import apply_style, korean_font

    apply_style()
    assert plt.rcParams["font.family"] == [korean_font()]
    # 한글 폰트에는 유니코드 마이너스가 없다. 음수 축이 깨진다.
    assert plt.rcParams["axes.unicode_minus"] is False


def test_설치된_한글_폰트를_고른다() -> None:
    """**png 는 서버에서 굽는다** (25세션).

    한 이름을 박아 두면 그 폰트가 없는 배포지(리눅스)에서만 한글이 네모로
    나온다 — 윈도우 개발 PC 에서는 드러나지 않는다. 설치된 것 중에서 고른다.
    """
    import matplotlib
    from matplotlib import font_manager

    from kwise.report import figures

    installed = {item.name for item in font_manager.fontManager.ttflist}
    assert figures.korean_font() in installed, "설치되지 않은 폰트를 골랐습니다."
    # 윈도우·macOS·리눅스 폰트가 모두 후보에 있어야 한 OS 에서만 되는 일이 없다.
    assert "Malgun Gothic" in figures.KOREAN_FONT_CANDIDATES  # 윈도우
    assert "AppleGothic" in figures.KOREAN_FONT_CANDIDATES  # macOS
    assert {"NanumGothic", "Noto Sans CJK KR"} <= set(figures.KOREAN_FONT_CANDIDATES)  # 리눅스
    assert matplotlib.get_backend()  # import 만으로 백엔드가 잡힌다


def test_한글_폰트가_없어도_멈추지_않는다(caplog: pytest.LogCaptureFixture) -> None:
    """**그림은 나와야 한다.** 한글이 깨지더라도 산출물 생성이 멈추면 더 나쁘다."""
    import logging

    from kwise.report import figures

    original = figures.KOREAN_FONT_CANDIDATES
    figures.korean_font.cache_clear()
    figures.KOREAN_FONT_CANDIDATES = ("있을 리 없는 폰트",)
    try:
        with caplog.at_level(logging.WARNING, logger="kwise.report.figures"):
            picked = figures.korean_font()
    finally:
        figures.KOREAN_FONT_CANDIDATES = original
        figures.korean_font.cache_clear()

    assert picked == figures.FALLBACK_FONT
    assert any("한글 폰트를 찾지 못해" in record.message for record in caplog.records)


def test_그림이_png_바이트다(sample_diagnosis: Diagnosis) -> None:
    from kwise.report import figures

    for png in (
        figures.hourly_profile_png(sample_diagnosis.peak),
        figures.monthly_peak_png(sample_diagnosis.peak),
        figures.top_hour_png(sample_diagnosis.peak),
    ):
        assert png.startswith(b"\x89PNG"), "png 매직 바이트가 아닙니다."
        assert len(png) > 1_000


# ===================================================================== ③ 표


def test_표가_word_표_객체다(full_document: DocumentType) -> None:
    """**이미지로 넣으면 제안서에 복사해 쓸 수 없다** (10.5)."""
    assert len(full_document.tables) >= 6
    for table in full_document.tables:
        assert table.style is not None
        assert table.style.name == TABLE_STYLE
        assert len(table.rows) >= 2
        assert table.rows[0].cells[0].text  # 머리글이 비어 있지 않다


def test_표_머리글이_굵다(full_document: DocumentType) -> None:
    header = full_document.tables[0].rows[0]
    assert any(run.font.bold for cell in header.cells for p in cell.paragraphs for run in p.runs)


def test_조합_비교가_표로_나온다(full_document: DocumentType) -> None:
    """**「요금제」 열은 S112 5절에 붙었다** (⑱).

    조합이 조합 부하에서 선택요금을 다시 고르게 되면서 **조합마다 다를 수
    있다** — 열이 없으면 2단계가 권한 하나로 전부 낸 줄 읽는다.
    """
    table = _table_with_header(full_document, "조합", "요금제", "절감액", "투자비", "회수기간")
    assert len(table.rows) >= 2  # type: ignore[attr-defined]
    body = [cell.text for cell in table.rows[1].cells]  # type: ignore[attr-defined]
    assert body[0]
    assert body[1].startswith("선택") or body[1] == "전체시간"
    assert "원" in body[2]  # 단위를 값에 붙인다 (Word 표는 열 이름에 단위가 없다)


def test_금액에_단위가_붙는다(full_document: DocumentType) -> None:
    table = _table_with_header(full_document, "항목", "값")
    values = [row.cells[1].text for row in table.rows[1:]]  # type: ignore[attr-defined]
    assert any(value.endswith("원") for value in values)


def test_요금_구조_표의_기본_더하기_전력량이_합계다(
    sample_usage: UsageData, sample_report: object, tariff: TariffTable
) -> None:
    """**표가 스스로 산수로 맞아야 한다** (S124 · ②-40).

    「현재 요금 구조」 는 앞서 ``base_won``(역률요금을 뺀 값)을 적고 비중은
    그것을 ``total_won``(담은 값)으로 나눈 몫을 적었다 — 역률 85% 를 걸면
    기본 + 전력량이 합계보다 317,220원 모자라고 비중 합이 99.8% 가 된다.
    화면·PPT 는 앞서부터 역률요금을 담아 세고 있었다.

    **간주 92% 벌에서는 역률요금이 0원이라 값이 그대로다** — 이 못이 벌을
    따로 세우는 까닭이다.
    """
    from kwise.diagnose import ContractInfo, diagnose
    from kwise.tariff import BillingOptions, TariffSelection, calculate_bill

    selection = TariffSelection("general_b", "high_a", "I")
    options = BillingOptions(power_factor_pct=85.0, contract_kw=5_500.0)
    bill = calculate_bill(sample_usage, tariff, selection, quality=sample_report, options=options)
    diagnosis = diagnose(
        sample_usage,
        tariff,
        ContractInfo(selection, contract_kw=5_500.0),
        quality=sample_report,
        options=options,
    )
    assert bill.total_power_factor_won > 0, "역률요금이 서는 벌이어야 한다"

    document = build_document(DocumentSections(usage=sample_usage, bill=bill, diagnosis=diagnosis))
    table = _table_with_header(document, "구분", "금액·비중")
    cells = {
        row.cells[0].text: row.cells[1].text  # type: ignore[attr-defined]
        for row in table.rows[1:]  # type: ignore[attr-defined]
    }

    def _won_of(label: str) -> int:
        return int(cells[label].split("원")[0].replace(",", ""))

    def _share_of(label: str) -> float:
        return float(cells[label].split("(")[1].rstrip("%)"))

    # 두 칸이 각자 천원 절사되므로 합계와 1,000원까지 어긋난다.
    total = int(
        money.won(diagnosis.structure.total_won, reason="—").removesuffix("원").replace(",", "")
    )
    assert abs(_won_of("기본요금") + _won_of("전력량요금") - total) <= 1_000
    # 소수 한 자리로 적으므로 0.1%p 까지.
    assert abs(_share_of("기본요금") + _share_of("전력량요금") - 100.0) <= 0.1


def test_요금_구조_표가_초과사용부가금_칸을_세운다(
    sample_usage: UsageData, sample_report: QualityReport, tariff: TariffTable
) -> None:
    """**부가금이 서는 벌에서 표가 합계에 못 미쳤다** (S127 2절 · ②-32).

    화면과 PPT 는 109세션부터 부가금 칸을 세우는데 **Word 표만 안 세웠다** —
    분모 ``total_won`` 은 부가금을 담으므로 기본 + 전력량이 합계에 못 미친다.
    위 못과 같은 병인데 **역률이 아니라 부가금 쪽**이고, 그래서 벌을 따로
    세운다: 계약전력이 관측 최대 위면 부가금이 0원이라 칸이 아예 안 생긴다.

    **계약전력 4,000 kW 는 덱 벌 `large-b-short` 와 같은 조건이다** — 저장소에
    부가금이 서는 벌이 하나도 없어 이 갈래가 한 번도 그려진 적이 없었다.
    """
    from kwise.diagnose import ContractInfo, diagnose
    from kwise.tariff import BillingOptions, TariffSelection, calculate_bill

    selection = TariffSelection("general_b", "high_a", "I")
    options = BillingOptions(contract_kw=4_000.0)
    bill = calculate_bill(sample_usage, tariff, selection, quality=sample_report, options=options)
    diagnosis = diagnose(
        sample_usage,
        tariff,
        ContractInfo(selection, contract_kw=4_000.0),
        quality=sample_report,
        options=options,
    )
    assert bill.total_excess_won > 0, "초과사용부가금이 서는 벌이어야 한다"

    document = build_document(DocumentSections(usage=sample_usage, bill=bill, diagnosis=diagnosis))
    table = _table_with_header(document, "구분", "금액·비중")
    cells = {
        row.cells[0].text: row.cells[1].text
        for row in table.rows[1:]  # type: ignore[attr-defined]
    }
    assert "초과사용부가금" in cells

    def _share_of(label: str) -> float:
        return float(cells[label].split("(")[1].rstrip("%)"))

    labels = ["기본요금", "전력량요금", "초과사용부가금"]
    assert abs(sum(_share_of(label) for label in labels) - 100.0) <= 0.1


# ===================================================================== ④ 미산출 사유


def test_미산출_항목에_사유가_들어간다(sample_usage: UsageData, sample_bill: BillingResult) -> None:
    """**빈칸도 0원도 아니다** (7.5). 0원은 '공짜' 로 읽힌다."""
    unpriced = SolarPoint(
        capacity_kwp=960.0,
        generation_kwh=1_125_000.0,
        self_consumed_kwh=1_000_000.0,
        surplus_kwh=125_000.0,
        self_consumption_ratio=0.89,
        billing_demand_kw=5_198.0,
        base_saving_won=1_000_000.0,
        energy_saving_won=2_000_000.0,
        total_saving_won=3_000_000.0,
        annual_saving_won=3_000_000.0,
        investment_won=None,  # 단가 미입력
        payback_years=None,
        power_factor_after_pct=92.0,
        power_factor_extra_won=0.0,
    )
    entry = measure_entries(solar=unpriced, solar_unpriced_reason="태양광 단가 미입력")[0]
    assert entry.investment != ""
    assert "0원" not in entry.investment
    assert "미산출" in entry.investment
    assert "단가" in entry.investment
    assert "미산출" in entry.payback


def test_보고서_본문에도_사유가_실린다(
    sample_usage: UsageData, sample_bill: BillingResult, sample_diagnosis: Diagnosis
) -> None:
    unpriced = SolarPoint(
        capacity_kwp=100.0,
        generation_kwh=1.0,
        self_consumed_kwh=1.0,
        surplus_kwh=0.0,
        self_consumption_ratio=1.0,
        billing_demand_kw=1.0,
        base_saving_won=0.0,
        energy_saving_won=0.0,
        total_saving_won=0.0,
        annual_saving_won=0.0,
        investment_won=None,
        payback_years=None,
        power_factor_after_pct=92.0,
        power_factor_extra_won=0.0,
    )
    document = build_document(
        DocumentSections(
            usage=sample_usage,
            bill=sample_bill,
            diagnosis=sample_diagnosis,
            measures=measure_entries(
                solar=unpriced, solar_unpriced_reason="단가를 넣지 않았습니다"
            ),
        )
    )
    assert "단가를 넣지 않았습니다" in _all_text(document)


def test_수단_항목이_모두_같은_틀이다(entries: tuple[object, ...]) -> None:
    for entry in entries:
        assert entry.conclusion  # type: ignore[attr-defined]
        assert entry.saving  # type: ignore[attr-defined]
        assert entry.investment  # type: ignore[attr-defined]
        assert entry.payback  # type: ignore[attr-defined]
        assert entry.certainty  # type: ignore[attr-defined]


def test_수단이_7장_순서로_나온다(
    sample_switch: TariffSwitchResult, sample_usage: UsageData, sample_bill: BillingResult
) -> None:
    from kwise.measures import evaluate_contract_adjustment, evaluate_power_factor
    from kwise.tariff import load_tariff

    table = load_tariff()
    power_factor: PowerFactorResult = evaluate_power_factor(
        sample_usage, table, sample_bill.selection, baseline=sample_bill
    )
    contract = evaluate_contract_adjustment(sample_usage, sample_bill, contract_kw=5_800.0)
    # 일부러 뒤죽박죽으로 넘긴다 — 순서는 목록이 정한다.
    built = measure_entries(power_factor=power_factor, switch=sample_switch, contract=contract)
    assert [item.kind.number for item in built] == ["7.1", "7.2", "7.4"]


def test_켜지_않은_수단은_보고서에_없다(entries: tuple[object, ...]) -> None:
    """**'보지 않은 것' 이 '검토했더니 이만큼' 으로 둔갑하면 안 된다.**"""
    numbers = {item.kind.number for item in entries}  # type: ignore[attr-defined]
    assert "7.5" not in numbers
    assert "7.6" not in numbers


# ===================================================================== ⑤ 검토 범위


def test_검토_범위가_마지막_장에_있다(full_document: DocumentType) -> None:
    table = _table_with_header(full_document, "구분", "수단")
    rows = {row.cells[0].text: row.cells[1].text for row in table.rows[1:]}  # type: ignore[attr-defined]
    assert rows["검토함"] == "7.1 선택요금 전환, 7.2 계약전력 조정"
    assert "7.5 태양광" in rows["미검토"]
    assert "7.6 ESS" in rows["미검토"]


def test_수단이_없으면_전부_미검토다(diagnosis_only: DocumentType) -> None:
    table = _table_with_header(diagnosis_only, "구분", "수단")
    rows = {row.cells[0].text: row.cells[1].text for row in table.rows[1:]}  # type: ignore[attr-defined]
    assert rows["검토함"] == "없음"
    for kind in MEASURE_CATALOG:
        assert kind.title in rows["미검토"]


def test_검토_범위를_넘겨받으면_그것을_쓴다(
    sample_usage: UsageData, sample_bill: BillingResult
) -> None:
    """화면의 :func:`review_scope` 결과를 그대로 싣는다."""
    sections = DocumentSections(
        usage=sample_usage,
        bill=sample_bill,
        reviewed_labels=("7.3 경제성DR",),
        skipped_labels=("7.1 선택요금 전환",),
    )
    assert sections.scope() == (("7.3 경제성DR",), ("7.1 선택요금 전환",))


def test_한계와_추적성이_마지막_장에_있다(full_document: DocumentType) -> None:
    text = _all_text(full_document)
    assert "기후환경요금" in text  # 미포함 요금요소 (5.1)
    assert "한 번의 초과가 12개월간 적용됩니다" in text  # 계약전력 경고 (9.4)
    assert "인증·신고용 산출물이 아닙니다" in text  # 알려진 한계 (부록 D)
    assert "적용 요금표:" in text  # 추적성 (5.8)
    assert "Open-Meteo" in text  # 출처


# ===================================================================== 감도


def test_감도를_범위로_적는다(
    sample_usage: UsageData, sample_bill: BillingResult, sample_diagnosis: Diagnosis
) -> None:
    """**3열 나열하지 않는다** (9.2)."""
    ranges = (
        SensitivityRange(
            "기본요금 절감액(원)", "원", 31_520_000, 28_970_000, 32_660_000, "첨예형", "평탄형"
        ),
    )
    document = build_document(
        DocumentSections(
            usage=sample_usage,
            bill=sample_bill,
            diagnosis=sample_diagnosis,
            comparison=None,
            sensitivity=ranges,
        )
    )
    text = _all_text(document)
    # 감도는 조합 비교 장에 붙는다. 조합이 없으면 실리지 않는다.
    assert "평탄형" not in text or "범위" in text


def test_감도_범위가_조합_장에_실린다(
    sample_usage: UsageData,
    sample_bill: BillingResult,
    sample_diagnosis: Diagnosis,
    sample_comparison: ComparisonResult,
) -> None:
    ranges = (
        SensitivityRange(
            "기본요금 절감액(원)", "원", 31_520_000, 28_970_000, 32_660_000, "첨예형", "평탄형"
        ),
    )
    document = build_document(
        DocumentSections(
            usage=sample_usage,
            bill=sample_bill,
            diagnosis=sample_diagnosis,
            comparison=sample_comparison,
            sensitivity=ranges,
        )
    )
    table = _table_with_header(document, "지표", "기준값과 범위")
    cell = table.rows[1].cells[1].text  # type: ignore[attr-defined]
    assert "범위" in cell
    assert "~" in cell
    # 시나리오 이름이 좋고 나쁨을 뜻하지 않는다는 안내가 함께 간다.
    assert "시나리오 이름" in _all_text(document)


# ===================================================================== ⑥ 한글


def test_한글이_깨지지_않는다(tmp_path: Path, full_sections: DocumentSections) -> None:
    """저장했다 다시 읽어도 한글이 그대로여야 한다."""
    path = export_document(full_sections, output_dir=tmp_path)
    text = _all_text(ReadDocument(str(path)))
    for expected in (
        DOCUMENT_TITLE,
        "전력 비용 진단 보고서",
        "요금적용전력",
        "검토 범위와 한계",
        "본사 사옥",
        "선택요금 전환",
    ):
        assert expected in text
    assert "?" not in text.replace("?", "", 0)  # 물음표 치환이 일어나지 않았다


def test_본문_글꼴이_한글_글꼴이다(full_document: DocumentType) -> None:
    """**동아시아 글꼴을 따로 지정해야** Word 가 한글에 쓴다."""
    from docx.oxml.ns import qn

    from kwise.report.document import KOREAN_FONT

    style = full_document.styles["Normal"]
    assert style.font.name == KOREAN_FONT
    assert style.element.rPr.rFonts.get(qn("w:eastAsia")) == KOREAN_FONT


# ===================================================================== 파일·바이트


def test_파일명에_날짜와_시각이_붙는다() -> None:
    """Word 가 파일을 열고 있으면 덮어쓰기가 실패한다."""
    path = document_path(Path("output"), now=dt.datetime(2026, 8, 11, 14, 30))
    assert path.name == "kwise_report_20260811_1430.docx"


def test_바이트로_받으면_디스크에_남지_않는다(
    tmp_path: Path, full_sections: DocumentSections
) -> None:
    payload, filename = document_bytes(full_sections, now=dt.datetime(2026, 8, 11, 9, 5))
    assert payload[:2] == b"PK"  # docx 는 zip 이다
    assert filename == "kwise_report_20260811_0905.docx"
    assert list(tmp_path.iterdir()) == []


def test_저장_경로에_폴더가_없으면_만든다(tmp_path: Path, full_sections: DocumentSections) -> None:
    target = tmp_path / "없던폴더"
    path = export_document(full_sections, output_dir=target)
    assert path.is_file()


# ===================================================================== 목록 동기화


def test_화면과_보고서가_같은_수단_목록을_본다() -> None:
    """각자 목록을 들고 있으면 한쪽만 고쳤을 때 조용히 어긋난다."""
    from kwise.ui.spec import MEASURES

    assert [item.kind for item in MEASURES] == list(MEASURE_CATALOG)
    # **7.7 잉여 활용은 41세션에 빠졌다** — 개선안이 아니라 태양광의 결과다.
    # 번호는 밀지 않는다 (:mod:`kwise.measures.catalog`).
    assert [item.number for item in MEASURE_CATALOG] == [
        "7.1",
        "7.2",
        "7.3",
        "7.4",
        "7.5",
        "7.6",
    ]


# ==================================================== 계산 근거와 부록 (22세션)


def test_계산_근거가_화면_Excel_Word_에서_같다(
    sample_usage: UsageData, sample_bill: BillingResult, sample_switch: TariffSwitchResult
) -> None:
    """**만드는 자리가 하나다** (22세션 2절).

    화면 카드·Excel 부록 A·Word 부록 A 가 같은 :class:`Worksheet` 를 쓴다.
    각자 표를 만들면 숫자가 갈라지고, 갈라진 것은 나란히 놓고서야 드러난다.
    """
    from kwise.report.appendix import worksheet_frame
    from kwise.report.document import DocumentSections, build_document
    from kwise.report.excel import ReportSections, build_sheets
    from kwise.report.worksheet import COLUMNS, tariff_switch_worksheet

    sheet = tariff_switch_worksheet(sample_switch)
    assert sheet, "선택요금 계산 근거가 비었습니다."
    screen = sheet.frame()
    assert list(screen.columns) == list(COLUMNS)
    # 산식과 대입값이 같은 줄에 있다.
    base = screen[screen["구분"] == "기본요금"].iloc[0]
    assert "kW ×" in base["산식"] and "원/kW" in base["산식"]
    assert base["값"].endswith("원")

    excel = build_sheets(ReportSections(usage=sample_usage, bill=sample_bill, worksheets=(sheet,)))[
        "부록 A 산출 근거"
    ]
    assert set(screen["값"]) <= set(excel["값"])
    assert list(excel.columns) == ["수단", *COLUMNS]
    assert worksheet_frame((sheet,)).equals(excel)

    document = build_document(
        DocumentSections(usage=sample_usage, bill=sample_bill, worksheets=(sheet,))
    )
    cells = {cell.text for table in document.tables for row in table.rows for cell in row.cells}
    for value in screen["값"]:
        if value:
            assert value in cells, value


def test_부록_B_는_기준_데이터에서_생성된다() -> None:
    """**손으로 옮겨 적지 않는다** (요구사항서 12장 · 22세션 3절)."""
    from kwise.report.appendix import reference_rows
    from kwise.rules import assumptions, rules

    rows = reference_rows()
    labels = {row[1] for row in rows}
    kinds = {row[0] for row in rows}
    assert kinds == {"법령 유래", "판단값"}
    assert len(rows) == len(rules().item_keys()) + len(assumptions().item_keys())
    for key in ("dr.market_hours", "power_factor.lagging_standard_pct"):
        assert rules()[key].label in labels, key
    # 근거 조문과 확인일이 함께 실린다 — 값만 있으면 출처를 되짚을 수 없다.
    statutory = [row for row in rows if row[0] == "법령 유래"]
    assert all(row[3] and row[3] != "—" for row in statutory)


def test_부록_C_가_한계와_참고를_한_곳에_모은다(
    sample_bill: BillingResult, sample_diagnosis: Diagnosis
) -> None:
    """**같은 말이 두 곳에 있으면 안 된다** (22세션 3절)."""
    from kwise.notices import Severity, texts
    from kwise.report.appendix import known_limits
    from kwise.report.notices import KNOWN_LIMITS

    lines = known_limits(sample_bill.notices, sample_diagnosis.notices)
    assert set(KNOWN_LIMITS) <= set(lines), "부록 D 목록이 빠졌습니다."
    for line in texts(sample_bill.notices, Severity.INFO):
        assert any(line[:30] in item for item in lines), line
    # 앞 30자가 같은 줄이 두 번 실리지 않는다.
    heads = [line[:30] for line in lines]
    assert len(heads) == len(set(heads))


def test_Word_의_1단계_결론과_7_2_결론이_같은_문장이다(tmp_path: Path, tariff: TariffTable) -> None:
    """**한 산출물 안에서 두 장이 반대로 말했다** (100세션 4절).

    종별을 넘는 벌에서 7.2 는 「299 kW 로 낮추면 종별이 바뀐다」 라고 적는데
    5.5 는 **「계약전력 400 kW 는 적정합니다」** 라고 적고 있었다. 5.5 가 하한
    한 줄만 보는 1단계 판정 위에 자기 문장을 따로 짓고 있었기 때문이다.
    이제 두 장이 :func:`_contract_conclusion` 하나를 쓴다.

    최대수요 200 kW · 계약 400 kW 의 을 고객이다. **하한 120 kW 는 진다** —
    같은 을 안에서는 낮출 이유가 없는데, 문턱 300 kW 아래인 299 kW 로 내리면
    일반용전력(갑)Ⅱ 로 넘어가 요금 전체가 준다 (99세션이 연 띠다).
    """
    from kwise.diagnose import ContractInfo, diagnose
    from kwise.io import load_usage
    from kwise.measures import evaluate_contract_adjustment
    from kwise.tariff import BillingOptions, TariffSelection, calculate_bill
    from tests._synthetic import make_labels, month_dates, write_csv

    rows = [
        (label, 50.0)  # 15분 50 kWh = 200 kW
        for date in month_dates(2024, 3)
        for label in make_labels(date)
    ]
    usage = load_usage(write_csv(tmp_path / "flat.csv", rows))
    selection = TariffSelection("general_b", "high_a", "I")
    options = BillingOptions(contract_kw=400.0)
    bill = calculate_bill(usage, tariff, selection, options=options)
    adjustment = evaluate_contract_adjustment(
        usage, bill, contract_kw=400.0, table=tariff, options=options
    )
    assert not adjustment.floor_binding, "하한이 이기면 결론이 하한 갈래로 간다."
    assert adjustment.crosses_type, "이 벌이 종별을 안 넘으면 시험이 뜻을 잃는다."

    document = build_document(
        DocumentSections(
            usage=usage,
            bill=bill,
            diagnosis=diagnose(
                usage, tariff, ContractInfo(selection, contract_kw=400.0), options=options
            ),
            measures=measure_entries(contract=adjustment),
        )
    )
    conclusions = [
        para.text.strip()
        for para in document.paragraphs
        if "일반용전력(갑)Ⅱ 로 바뀌어" in para.text
    ]
    # 5.5(1단계 적정성)와 7.2(계약전력 조정) 둘이다. 글자까지 같아야 한다.
    assert len(conclusions) == 2, conclusions
    assert conclusions[0] == conclusions[1]
    assert "적정합니다" not in _all_text(document).split("검토 범위")[0]


def test_Word_7_2_주의사항에_같은_경고가_두_번_서지_않는다(
    tmp_path: Path, tariff: TariffTable
) -> None:
    """**같은 문장이 잇달아 두 번 섰다** (84세션 · 100세션 4절 — 같은 결함 둘).

    `report\\notices.py` 의 ``CONTRACT_CHANGE_WARNING`` 과
    `measures\\contract.py` 의 ``MARGIN_NOTICE`` 가 **글자까지 같은 사본**이라,
    7.2 주의사항이 앞에 세운 한 줄과 ``contract.notices`` 에서 온 한 줄을
    **둘 다** 실었다. 화면은 같은 문자열을 걸러 내고 있었고 Word 만 안 걸렀다.

    **글자를 세는 못이 아니라 실물을 지어 세는 못이다** — 문서를 실제로 짓고
    7.2 절 안의 문단만 골라 센다. 102세션에 고치기 전 값은 **2회**였다.
    **사본 둘을 합치는 것은 이 못의 몫이 아니다** — 뿌리는 미해결에 남아 있다.
    """
    from kwise.diagnose import ContractInfo, diagnose
    from kwise.io import load_usage
    from kwise.measures import evaluate_contract_adjustment
    from kwise.notices import texts
    from kwise.report import CONTRACT_CHANGE_WARNING
    from kwise.tariff import BillingOptions, TariffSelection, calculate_bill
    from tests._synthetic import make_labels, month_dates, write_csv

    rows = [
        (label, 50.0)  # 15분 50 kWh = 200 kW
        for date in month_dates(2024, 3)
        for label in make_labels(date)
    ]
    usage = load_usage(write_csv(tmp_path / "flat.csv", rows))
    selection = TariffSelection("general_b", "high_a", "I")
    options = BillingOptions(contract_kw=400.0)
    bill = calculate_bill(usage, tariff, selection, options=options)
    adjustment = evaluate_contract_adjustment(
        usage, bill, contract_kw=400.0, table=tariff, options=options
    )
    # **낮출 자리가 있어야 그 경고가 실린다** — 없으면 시험이 뜻을 잃는다.
    assert adjustment.reducible
    assert CONTRACT_CHANGE_WARNING in texts(adjustment.notices)

    document = build_document(
        DocumentSections(
            usage=usage,
            bill=bill,
            diagnosis=diagnose(
                usage, tariff, ContractInfo(selection, contract_kw=400.0), options=options
            ),
            measures=measure_entries(contract=adjustment),
        )
    )
    paragraphs = list(document.paragraphs)
    heads = [
        index
        for index, para in enumerate(paragraphs)
        if _style_name(para).startswith("Heading")
    ]
    start = next(index for index in heads if "7.2 계약전력 조정" in paragraphs[index].text)
    end = next((index for index in heads if index > start), len(paragraphs))
    block = [para.text.strip() for para in paragraphs[start:end]]

    assert block.count(CONTRACT_CHANGE_WARNING) == 1, block
    assert "주의사항" in block, "주의사항 자리 자체가 사라지면 안 된다."
