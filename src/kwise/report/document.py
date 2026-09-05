"""Word 보고서 (요구사항서 10.5).

Excel 여덟 시트는 **분석자용 데이터**다. 의사결정자가 읽을 문서가 따로 필요하다.

    각 절은 **결론부터** 쓴다. 데이터를 앞에 놓지 않는다.
    표는 **Word 표 객체**로 만든다. 이미지로 넣으면 제안서에 복사해 쓸 수 없다.
    차트만 png 다 (:mod:`kwise.report.figures`).
    감도는 **범위**로 적는다. 3열 나열은 하지 않는다 (9.2).

**수단을 하나도 켜지 않아도 보고서가 나온다.** 진단만 보고 받아 가는 것이 정상
경로다 (8세션에서 같은 종류의 결함을 Excel 에서 잡았다). 그때는 3·4장을 빼고
장 번호를 당기며, 무엇을 보지 않았는지는 마지막 장의 「검토 범위」가 밝힌다.

파일명에 날짜·시각 접미사를 붙인다 — Word 가 파일을 열고 있으면 덮어쓰기가
실패한다 (Excel 과 같은 이유).
"""

from __future__ import annotations

import datetime as dt
import io
from collections.abc import Callable, Iterable, Sequence
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from kwise.compare import SCENARIO_NAME_CAVEAT, ComparisonResult, SensitivityRange
from kwise.diagnose import ContractAdequacy, Diagnosis
from kwise.diagnose.dr import DR_OFF_DAYS_FACT, DrProfile
from kwise.io import UsageData
from kwise.measures import (
    MEASURE_CATALOG,
    NO_SAVING,
    NOT_VIABLE_CONCLUSION,
    Certainty,
    ContractAdjustment,
    DemandResponseResult,
    EssOptimum,
    EssResult,
    EssTargetCurve,
    MeasureKind,
    PowerFactorResult,
    SolarPoint,
    SurplusResult,
    SurplusScenario,
    TariffSwitchResult,
    annualize,
    measure_kind,
    payback_text,
    spec_mark_note,
)
from kwise.measures import surplus as surplus_module
from kwise.notices import Notice, Severity, report_appendix, report_body, texts
from kwise.report import figures, frames, narrative
from kwise.report.appendix import APPENDIX_TITLES, AppendixData, known_limits, reference_rows
from kwise.report.days import RepresentativeDay
from kwise.report.notices import (
    CONTRACT_CHANGE_WARNING,
    DATA_SOURCES,
    NOT_INCLUDED_NOTICE,
    TRUNCATION_FOOTNOTE,
    UNPRICED_REASONS,
    format_won,
    plain_text,
)
from kwise.report.worksheet import COLUMNS, Worksheet
from kwise.tariff import BillingResult, TariffTable
from kwise.tariff.labels import option_label
from kwise.tariff.power_factor import lagging_standard_pct

__all__ = [
    "CHAPTER_COMPARISON",
    "CHAPTER_DIAGNOSIS",
    "CHAPTER_MEASURES",
    "CHAPTER_SCOPE",
    "CHAPTER_SUMMARY",
    "DEFAULT_OUTPUT_DIR",
    "DOCUMENT_TITLE",
    "KOREAN_FONT",
    "MEASURE_STRIP_FIGURE",
    "SURPLUS_CHOSEN_MARK",
    "SURPLUS_PAGE_NOTE",
    "SURPLUS_SCENARIO_HEADER",
    "TABLE_STYLE",
    "DocumentSections",
    "MeasureEntry",
    "MeasureFigure",
    "SurplusPage",
    "build_document",
    "document_bytes",
    "document_path",
    "export_document",
    "measure_entries",
    "surplus_page",
]

DOCUMENT_TITLE = "전력 비용 진단 보고서"

#: Word 본문 글꼴 **이름**.
#:
#: 차트 png(:mod:`kwise.report.figures`)와 달리 **여기서 글자를 그리지 않는다** —
#: 이름만 문서에 적고 그리는 것은 읽는 사람의 Word 다. 그 PC 에 없으면 Word 가
#: 알아서 대체하므로, 서버에 이 폰트가 깔려 있을 필요가 없다. 받는 사람이 대개
#: 한국 사무실이라 맑은 고딕으로 둔다.
KOREAN_FONT = "Malgun Gothic"
TABLE_STYLE = "Table Grid"
DEFAULT_OUTPUT_DIR = Path("output")

CHAPTER_SUMMARY = "요약"
CHAPTER_DIAGNOSIS = "진단"
CHAPTER_MEASURES = "개선 수단별 검토"
CHAPTER_COMPARISON = "조합 비교"
CHAPTER_SCOPE = "검토 범위와 한계"

_FIGURE_WIDTH = Inches(6.3)
_UNPRICED = "미산출"


# ===================================================================== 수단 한 항목


@dataclass(frozen=True)
class MeasureFigure:
    """수단 한 장에 나란히 놓을 그림 하나 (38세션 3절).

    **PPT 가 쓴다.** 화면은 역률·태양광·ESS 에 그림을 둘씩 그리는데 PPT 는
    하나씩만 실어, 「무엇을 보고 그렇게 판단했나」 가 슬라이드에서 빠져 있었다.
    차례가 곧 좌→우다.

    Word 는 :attr:`MeasureEntry.figure` 하나만 쓴다 — 문서는 절마다 세로로
    쌓이는 자리라 좌우로 나눌 칸이 없고, 36세션에 화면에서 감춘 산출물이다.
    """

    png: bytes
    caption: str


@dataclass(frozen=True)
class MeasureEntry:
    """3장의 수단 하나. **모든 수단을 같은 틀로 적는다.**

    결론 → 절감액 → 투자비 → 회수기간 → 주의사항.

    **확실성은 적지 않는다** (53세션 1-4). :attr:`certainty` 는 계산이 낸 값이라
    남겨 두되 어느 산출물도 싣지 않는다.

    금액 칸은 **문자열**이다. 산출하지 못한 항목에 빈칸이나 0원을 넣지 않고
    사유를 적기 위해서다 (7.5·7.6).
    """

    kind: MeasureKind
    conclusion: str
    saving: str
    investment: str
    payback: str
    certainty: str
    cautions: tuple[str, ...] = field(default=())
    notices: tuple[Notice, ...] = field(default=())
    """이 수단이 낸 안내 **원본**. 부록이 참고 등급만 골라 쓴다 (19세션 1절)."""
    figure: bytes | None = None
    """수단 차트 png (15세션). **화면과 같은 프레임을 본다** — 각자 만들면 어긋난다."""
    figure_caption: str = ""
    figures: tuple[MeasureFigure, ...] = field(default=())
    """PPT 가 좌우로 나란히 놓을 그림들 (38세션 3절). 비면 :attr:`figure` 한 장이다."""
    facts: tuple[tuple[str, str], ...] = field(default=())
    """**여지가 없다는 것을 보이는 지표** (39세션 4-2·4-3).

    결론이 「하향 여지 없음」·「잉여 0」 이면 실행 주의사항은 실을 자리가 아니다 —
    하지도 않을 일을 조심하라는 글이 결론보다 길어진다. 대신 **왜 없는지 보이는
    숫자**를 그 자리에 세운다. 화면 카드가 같은 판단으로 세운 지표들이다
    (31세션 2절·6절)."""
    actionable: bool = True
    """실행할 것이 있는가. **거짓이면 슬라이드가 주의사항을 싣지 않는다** (4-2)."""
    facts_first: bool = False
    """지표를 **결과보다 먼저** 세우는가 (53세션 6-1).

    계약전력 조정이 그렇다 — 「6,000 kW 를 5,823 kW 로」 는 **여유가 얼마나
    있는지를 보고 나서야** 판단할 수 있는 값인데, 절감액·투자비·회수기간이
    위에 서 있어 근거가 결과 뒤로 밀려 있었다. 참이면 :attr:`facts` 가 결론
    바로 아래에 서고 결과 지표가 그 아래로 내려간다."""
    has_saving: bool = True
    """절감액이 **0도 미산출도 아닌가.** 금액 문자열을 되파싱하지 않으려고
    숫자를 아는 자리에서 정해 둔다 — 부록이 실을 수단을 이것으로 고른다
    (39세션 5절)."""
    slide_note: str = ""
    """슬라이드 맨 아래 작은 글씨 한 줄 (39세션). **숫자가 왜 달라졌는지**를
    되짚어야 하는 사실이 여기 온다 — 사람이 쉬는 날을 빼면 감축 가능량이 다시
    계산되므로 그 사실이 산출물에 남아야 한다 (29세션)."""
    saving_annual: str = ""
    """**12개월 환산 한 값만** (39세션 2-2). 슬라이드가 쓴다 — 비면 :attr:`saving`.

    Word 는 :attr:`saving` 을 그대로 쓴다. 문서는 관측 기간 값과 환산값을 나란히
    두고 대조하는 자리라 두 값이 함께 있어야 한다."""
    spec_table: tuple[tuple[str, ...], ...] = field(default=())
    """산출물에 **폭 전체로** 놓을 표. 머리글이 첫 줄이다 (46세션).

    ESS 가 쓴다 — 회수기간 곡선을 목표별 사양 표로 바꾸면서 생겼다. PPT 는
    **반 칸에 놓지 않는다**: 열이 아홉이라 좌우로 나누면 글자가 읽히지 않는다.
    표가 있으면 그림은 표 아래로 내려간다. Word 는 절 안에 그대로 쌓는다."""
    spec_caption: str = ""
    """표 아래 한 줄. 화면 캡션과 **같은 문장이다.**"""
    spec_note: str = ""
    """표에 붙은 **표식의 뜻** 한 줄 (53세션 4-13).

    대형 자료의 사양 표에 「마진 미달」 이 셋인데 덱 어디에도 그 뜻이 없었다.
    :func:`~kwise.measures.spec_mark_note` 가 **표에 실제로 붙은 표식만** 골라
    적는다 — 셋을 늘 깔면 없는 표식을 설명하게 된다."""

    @property
    def slide_saving(self) -> str:
        """슬라이드에 적을 절감액. **한 칸에 한 값이다.**"""
        return self.saving_annual or self.saving

    @property
    def slide_figures(self) -> tuple[MeasureFigure, ...]:
        """슬라이드에 실을 그림 차례. **둘을 주지 않았으면 있는 하나를 쓴다.**"""
        if self.figures:
            return self.figures
        if self.figure is None:
            return ()
        return (MeasureFigure(self.figure, self.figure_caption),)

    @property
    def title(self) -> str:
        return self.kind.title


def body_lines(notices: tuple[Notice, ...]) -> tuple[str, ...]:
    """보고서 **본문**에 실을 줄 — 차단·주의·근거 (19세션 1절).

    **근거가 본문으로 온다.** 화면에서는 툴팁으로 접혀 있던 산식·출처·계수가
    여기서는 펼쳐져야 한다 — 보고서는 나중에 혼자 읽는 문서이고, 그때 물을 것이
    "이 숫자가 어디서 나왔나" 이기 때문이다. 참고는 5장 부록으로 간다.
    """
    return texts(report_body(notices))


#: 수단 한 장에 그림을 **둘 나란히** 놓을 때 굽는 크기 (38세션 3절).
#:
#: 반 칸(약 6in)에 세로 3.2in 남짓이라, 기본 비율(9:3.6)로 구우면 폭이 먼저 차서
#: 슬라이드에서 다시 줄어들고 축 눈금이 뭉개진다. 칸의 가로세로에서 나온 값이다.
MEASURE_PAIR_FIGURE = (6.0, 3.2)

#: 그림이 **하나뿐인** 수단 장의 크기 (38세션 4절). 폭을 다 쓰는 칸이다.
#:
#: 기본 비율(9:3.6)로 구우면 슬라이드에서 높이가 먼저 차서 폭의 3분의 2만 쓴다.
MEASURE_FULL_FIGURE = (12.0, 3.5)

#: 요금제 전환은 **위·아래 두 칸짜리** 한 장이다 (그룹 막대 + 현행 대비 차액).
#: 세로가 더 필요해 위 값을 그대로 쓸 수 없다.
TARIFF_FIGURE = (12.0, 4.4)

#: **표 아래로 내려앉는 그림**의 크기 (53세션 2절). ESS 가 쓴다.
#:
#: 목표별 사양 표가 위를 차지하고 나면 아래에 1in 남짓만 남는다. 그런데 그림은
#: ``MEASURE_PAIR_FIGURE`` 비율(6:3.2)로 구워져 **높이가 먼저 차** 폭의 7%만
#: 쓰고 0.85 × 0.41in 로 앉았다 — 손톱만 한 차트다. 납작하게 구우면 같은 높이로
#: 폭을 다 쓴다 (4·5장의 :data:`~kwise.report.slides.FULL_FIGURE` 와 같은 이치).
MEASURE_STRIP_FIGURE = (12.0, 2.2)

#: 그림 캡션. **한 자리에 둔다** — 같은 그림이 :attr:`MeasureEntry.figure` 와
#: :attr:`MeasureEntry.figures` 두 자리에 실리므로, 문구를 두 벌로 적으면 한쪽만
#: 고쳐진다.
_PF_TRIANGLE_CAPTION = "전력삼각형 — 각이 좁아질수록 역률이 좋습니다."
_PF_DAY_CAPTION = "대표일 부하 — 주황 점이 역률을 판정하는 주간 구간입니다."
_SOLAR_ANNUAL_CAPTION = "일별 발전량 — 여름에 높고 겨울에 낮습니다."
_SOLAR_DAY_CAPTION = "대표일의 부하 — 두 선 사이가 태양광으로 줄어든 몫입니다."
_ESS_DAY_CAPTION = "대표일의 부하 — 두 선 사이가 ESS 로 깎은 몫입니다."


def _contract_saving(contract: ContractAdjustment, value: float | None) -> str:
    """계약전력 조정의 절감액 칸 — **셋으로 갈린다** (48세션 · 83세션).

    미산출  하한 규정을 모른다
    없음    하한이 요금적용전력에 걸리지 않아 줄 것이 없다
    금액    실제로 준다

    **「없음」 에는 사유를 붙이지 않는다** (83세션). 같은 장의 각주가
    ``contract.floor_not_binding`` 로 이미 그 까닭을 적는다 — 사유를 함께
    적으면 「하한 30% 적용」 이라는 각주가 하나 더 서서, **걸리지도 않은 하한을
    적용했다고 읽힌다.**
    """
    if value is None:
        return f"{_UNPRICED} — {contract.saving_basis}"
    if contract.no_saving:
        return NO_SAVING
    return _won(value)


def _contract_adequacy_saving(adequacy: ContractAdequacy) -> str:
    """1단계 계약전력 적정성의 절감액 칸. 수단 쪽과 **같은 세 갈래다.**

    **「낮출 자리가 있다」 로 가른다** (100세션). ``floor_binding`` 으로 가르면
    하한이 지면서 종별을 넘는 판에서 절감액이 있는데도 「없음」 이 나간다.
    """
    if adequacy.saving_won is None:
        return f"{_UNPRICED} — {adequacy.saving_basis}"
    if not adequacy.reducible:
        return NO_SAVING
    return _won(adequacy.saving_won)


#: ESS 결론의 앞부분. **「회수기간이 가장 짧은 지점」 은 그대로 둔다** —
#: 자료 여덟에서 확인했다 (52세션 2절).
_ESS_TARGET_SENTENCE = "회수기간이 가장 짧은 지점을 목표로 잡았습니다."


def _ess_conclusion(ess: EssResult) -> str:
    """ESS 결론 — **성립하지 않으면 그 말을 붙인다** (53세션 4-13).

    「회수기간이 가장 짧은 지점을 목표로 잡았습니다」 는 맞다. 그런데 대형
    자료의 31.7년은 **배터리 보증 수명의 세 배**이고, 그 사실이 결론 어디에도
    없었다 — 「가장 짧다」 만 읽으면 성립하는 안으로 보인다.

    **문장을 새로 짓지 않는다.** 계산이 이미 낸 경고
    (``ess.payback_over_warranty``)를 그대로 옮긴다 — 수명 값도 판정도 거기
    하나가 쥔다. 슬라이드는 주의사항 표를 그리지 않으므로(그림이 그 자리를
    쓴다) 이 경고가 덱에서 통째로 빠지고 있었다.
    """
    head = (
        f"피크를 {ess.excess.target_kw:,.0f} kW 까지 낮추려면 ESS "
        f"{ess.power_kw:,.0f} kW / {ess.capacity_kwh:,.0f} kWh 가 필요합니다. "
        f"{_ESS_TARGET_SENTENCE}"
    )
    warranty = _notice_text(ess.notices, "ess.payback_over_warranty")
    return f"{head} {warranty}" if warranty else head


def _power_factor_conclusion(result: PowerFactorResult) -> str:
    """역률 결론 — **세 갈래다** (53세션 4-11).

    39세션까지는 「올리면 ○○원 줄어듭니다」 하나였다. 그런데 **현재 역률이
    기준에 못 미치면 그것은 감액이 아니라 추가요금 회피**다 — 이미 나가고 있는
    돈을 안 나가게 하는 것이라 성격이 다르다. 계산은 그 사실을 알고 있었다
    (:attr:`~kwise.measures.PowerFactorResult.is_penalty_removal`).

        추가요금   기준 미달. 없애는 쪽이 감액보다 금액이 크다
        감액       기준 이상에서 더 올린다
        여지 없음  목표가 현재보다 높지 않다 (이미 상한이거나 내리는 쪽)
    """
    current = f"{result.current_pct:,.0f}%"
    target = f"{result.target_pct:,.0f}%"
    if result.improvement_pct <= 0:
        return f"지상역률 {current} 에서 목표 {target} 로 올릴 여지가 없습니다."
    if result.is_penalty_removal:
        standard = f"{lagging_standard_pct():,.0f}%"
        return (
            f"지상역률 {current} 는 기준 {standard} 에 못 미쳐 기본요금에 추가요금이 "
            f"붙습니다. {target} 로 올리면 추가요금이 없어지고 감액을 받아 "
            f"{_won(result.saving_won)} 줄어듭니다."
        )
    return f"지상역률을 {current} → {target} 로 올리면 {_won(result.saving_won)} 줄어듭니다."


#: 계약전력 그림의 캡션 (53세션 6-3 · 83세션에 하한 선을 넣으며 고쳤다).
#: **「틈」 이라 부르지 않는다** — 틈은 기회를 암시하는데, 하한 아래의 틈은
#: 낮춰도 아무것도 주지 않는다.
_CONTRACT_HEADROOM_CAPTION = "붉은 선이 계약전력, 점선이 요금적용전력 하한입니다."


def _contract_facts(contract: ContractAdjustment) -> tuple[tuple[str, str], ...]:
    """판정을 가르는 수 — **화면 2단계와 같은 셋**이다 (83세션 7).

    하한이 이길 때만 목표 계약전력이 한 칸 더 선다. 「여유 %」 와 「하향 여지」
    는 뺐다 — 둘 다 여유율을 잣대로 삼는 값이라 판정과 어긋났다.
    """
    ratio = contract.contract_floor_ratio
    facts = [
        ("현재 계약전력", f"{contract.contract_kw:,.0f} kW"),
        (
            f"계약전력의 {ratio:.0%}" if ratio is not None else "하한",
            f"{contract.floor_kw:,.0f} kW" if contract.floor_kw is not None else "—",
        ),
        ("최대수요", f"{contract.demand_before_floor_kw:,.0f} kW"),
    ]
    if contract.target_contract_kw is not None:
        facts.append(("목표 계약전력", f"{contract.target_contract_kw:,.0f} kW"))
    return tuple(facts)


def _contract_conclusion(contract: ContractAdjustment) -> str:
    """계약전력 결론 — **네 갈래다** (53세션 4-9 · 83세션에 판정을 하한으로 옮겼다).

        초과 위약  계약전력을 넘은 구간이 있다. **상향을 검토할 자리다**
        종별이 바뀐다  문턱 아래 종별로 넘어가면 요금 전체가 준다 (99세션)
        하한이 걸린다  하한이 어느 달의 요금적용전력을 끌어올린다. 낮추면 그만큼 준다
        하한이 안 걸린다  어느 달도 안 걸리고 넘어갈 종별도 없어 낮춰도 줄지 않는다

    39세션까지는 「하향 여지 있음」 과 「적정」 둘이었고 여유율이 그 둘을 갈랐다.
    **여유율은 이 판정에 쓸 잣대가 아니다** — 이득이 있느냐는 하한이 정한다.

    **105세션에 두 갈래를 손봤다** (②-13). ㄱ 종별 갈래를 하한 갈래보다 **먼저**
    본다 — 둘 다 서는 판(하한이 걸리면서 문턱도 넘는 판)에서 절감액은 종별
    전환에서 오는데 하한 문장이 「그만큼 기본요금이 줄어듭니다」 라고 적고
    있었다. ㄴ 하한 문장이 **「최대수요보다 높아」 를 안 쓴다** — 연간 최대만
    보는 말이라, 굴림 창을 못 채운 초기 달에만 걸리는 판에서 거짓이 된다.
    **문구는 안 늘었다** — 갈래 넷이 그대로 넷이다.
    """
    if contract.over_contract_slots:
        return (
            f"계약전력 {contract.contract_kw:,.0f} kW 를 넘은 구간이 "
            f"{contract.over_contract_slots:,}건 있어 초과 위약 검토 대상입니다. "
            "하향이 아니라 상향을 검토해야 합니다."
        )
    if contract.target_contract_kw is not None and contract.crossed_label is not None:
        return (
            f"계약전력을 {contract.contract_kw:,.0f} → "
            f"{contract.target_contract_kw:,.0f} kW 로 낮추면 계약종별이 "
            f"{contract.crossed_label} 로 바뀌어 요금 전체가 줄어듭니다."
        )
    if contract.target_contract_kw is not None and contract.floor_bound_months:
        return (
            f"요금적용전력 하한 {contract.floor_kw:,.0f} kW 가 "
            f"{len(contract.floor_bound_months)}개 월의 요금적용전력을 끌어올리고 있어, "
            f"계약전력을 {contract.contract_kw:,.0f} → "
            f"{contract.target_contract_kw:,.0f} kW 로 낮추면 그만큼 기본요금이 줄어듭니다."
        )
    if contract.floor_kw is None:
        return f"계약전력 {contract.contract_kw:,.0f} kW 의 하한 비율을 알 수 없습니다."
    return (
        f"계약전력 {contract.contract_kw:,.0f} kW 는 이미 적정합니다. 하한 "
        f"{contract.floor_kw:,.0f} kW 가 어느 달의 요금적용전력에도 걸리지 않아 "
        "낮출 자리가 없습니다."
    )


def _shortest_discharge_hours(curve: EssTargetCurve) -> float:
    """곡선에서 가장 짧은 방전시간. 성립 한계와 나란히 놓아 격차를 보인다."""
    return min((item.discharge_hours for item in curve.points), default=0.0)


def _pair(*items: tuple[bytes | None, str]) -> tuple[MeasureFigure, ...]:
    """구운 그림들을 슬라이드 차례로 묶는다. **못 구운 것은 빠진다.**

    하나만 남으면 슬라이드가 그 하나를 크게 그린다 — 빈 칸을 남기지 않는다.
    """
    return tuple(MeasureFigure(png, caption) for png, caption in items if png is not None)


def _safe_figure(make: Callable[[], bytes], what: str = "이름 없는 그림") -> bytes | None:
    """그림을 굽되 **실패해도 보고서를 죽이지 않는다.**

    차트 하나 때문에 문서 전체를 잃는 편보다 그림 없이 표만 내는 편이 낫다
    (13세션에 3단계 화면에서 같은 종류의 결함을 겪었다).

    **조용히 삼키지는 않는다** (60세션 11절). 그림이 빠진 장을 사용자가 받아도
    알 길이 없으면 21세션이 걷어낸 「조용한 폴백」 과 같아진다 — 폴백은 두되
    :func:`~kwise.report.figures.note_figure_failure` 로 남긴다.

    Args:
        make: 그림을 굽는 일.
        what: 어느 수단의 어느 그림인지 — 기록에 그대로 실린다.
    """
    try:
        return make()
    except Exception as exc:
        figures.note_figure_failure(what, exc)
        return None


def _won(value: float | None, *, reason: str | None = None) -> str:
    """금액 한 칸. **모르면 빈칸이 아니라 사유다** (7.5).

    :func:`kwise.report.notices.format_won` 은 Excel 용이라 단위를 붙이지 않는다
    (열 이름이 ``(원)`` 을 달고 있다). Word 표는 열 이름에 단위가 없으므로
    값에 붙인다 — 숫자만 있는 칸은 읽는 사람이 단위를 되물어야 한다.
    """
    if value is None:
        return format_won(None) if reason is None else reason
    return f"{format_won(value)}원"


def _payback_text(years: float | None, investment_won: float | None) -> str:
    """회수기간 한 줄. **표시 상한을 넘으면 「>50년」 이다** (50세션 3-7)."""
    if investment_won is None:
        return f"{_UNPRICED} — 투자비 미입력"
    if not investment_won:
        return "즉시 (투자 없음)"
    return payback_text(years) if years is not None else f"{_UNPRICED} — 절감액 없음"


def _measure_saving(annual_won: float | None, period_won: float | None) -> str:
    if period_won is None:
        return _won(None)
    if annual_won is None:
        return _won(period_won)
    return f"{_won(period_won)} (12개월 환산 {_won(annual_won)})"


def _annual_saving(annual_won: float | None, period_won: float | None) -> str:
    """**12개월 환산 한 값만** (39세션 2-2).

    슬라이드는 한 칸에 두 값을 담지 않는다 — 「9,050,000원 (12개월 환산
    9,050,000원)」 은 같은 값을 두 번 적어 두 줄로 흐르고, 읽는 사람은 어느 쪽이
    답인지 되묻는다. 환산 기준이라는 사실은 **각주가 한 번** 적는다.
    """
    if annual_won is None and period_won is None:
        return _won(None)
    return _won(annual_won if annual_won is not None else period_won)


def _notice_text(notices: tuple[Notice, ...], fact: str) -> str:
    """사실 ID 로 안내 한 건을 꺼낸다. 없으면 빈 글이다.

    **문구를 새로 짓지 않는다** — 계산 모듈이 낸 글을 그대로 낸다 (31세션 0-2 의
    ``dropped_rows`` 와 같은 규약).
    """
    return next((item.text for item in notices if item.fact == fact), "")


def _dr_conclusion(result: DemandResponseResult, profile: DrProfile | None) -> str:
    """경제성DR 결론 (39세션 4-1).

    **0인 까닭이 함께 있어야 한다.** 거래 가능일과 저부하 평일만 적으면 왜 그
    날이 며칠뿐인지 알 수 없고, 결과가 0이면 「검토하지 않았다」 로 읽힌다 —
    저부하 판정은 **부하가 쉬는 날 수준까지 내려온 평일**을 세는 것이므로,
    그런 날이 없다는 사실이 곧 「추가로 줄일 여지가 없다」 는 답이다.

    **사람이 뺀 날은 이름으로 적는다** (29세션). 저부하로 잡혔다가 쉬는 날로
    판정되어 빠진 날이 있으면 숫자가 왜 달라졌는지 되짚을 수 있어야 한다.
    """
    head = narrative.dr_lead(profile)
    if not result.low_load_days:
        return head
    tail = (
        f" 등록 권장 {result.registered_capacity_kw:,.0f} kW, 연간 감축 가능량 "
        f"{result.annual_reducible_kwh:,.0f} kWh 입니다."
    )
    return head + tail


def _solar_conclusion(
    solar: SolarPoint,
    surplus_free_kwp: float | None,
    area_m2: float | None,
    *,
    surplus_page_follows: bool = False,
) -> str:
    """태양광 결론 한 줄 (39세션 3-1).

    **용량을 정한 근거가 함께 있어야 한다** — 얼마나 넓은 자리가 드는지,
    어디서부터 잉여가 생기는지가 그 근거다 (31세션 4-1 이 화면에 세운 값이다).
    「올리면」 은 우리끼리 쓰는 말이라 「설치하면」 으로 적는다.

    **한계를 넘겨 권할 때는 남는 양을 적는다** (51세션 3절). 소형 사무빌딩에서
    자가소비 한계가 40 kWp 인데 160 kWp 를 권하고 있었다 — **네 배다.** 그런데
    문장은 「40 kWp 까지는 전량 자가소비」 라고만 적어, 잉여가 생긴다는 사실도
    얼마나 생기는지도 말하지 않았다. **얼마가 남는지가 용량을 정하는 판단
    재료다.** 자리는 그대로 쓰고 말만 바꾼다 — 문구를 늘리지 않는다.
    """
    parts = [
        f"태양광 {solar.capacity_kwp:,.0f} kWp 를 설치하면 연 "
        f"{solar.generation_kwh:,.0f} kWh 를 발전해 "
        f"{_won(solar.total_saving_won)} 줄어듭니다."
    ]
    # **면적 상한과 자가소비 한계를 같은 층위로 나열하지 않는다** (53세션 4-12).
    # 51세션까지는 「(설치 면적 2,000 m² · 2,038 kWp 까지는 전량 자가소비)」 였는데,
    # **2,038 kWp 를 지을 수 있는 것처럼 읽힌다** — 지을 수 있는 것은 160 kWp 다.
    # 하나는 실제 상한이고 하나는 참고값이므로 문장을 나눈다.
    if area_m2:
        parts.append(f"설치 가능 면적 {area_m2:,.0f} m² 가 허용하는 상한입니다.")
    if surplus_free_kwp is not None and surplus_free_kwp > 0:
        if solar.capacity_kwp > surplus_free_kwp + 1e-9:
            tail = f"잉여 없이 지을 수 있는 것은 {surplus_free_kwp:,.0f} kWp 까지입니다."
            # **잉여량은 다음 장이 다룬다** (53세션 3절·4-12). 잉여 활용 장이
            # 생기지 않는 경우에만 여기서 얼마가 남는지 적는다 — 그 수는
            # 용량을 정하는 판단 재료라 어디엔가는 있어야 한다 (51세션 3절).
            if not surplus_page_follows:
                tail += f" 이 용량에서는 {solar.surplus_kwh:,.0f} kWh 가 남습니다."
            parts.append(tail)
        else:
            parts.append("이 용량에서는 발전량이 전부 자가소비됩니다.")
    return " ".join(parts)


# ===================================================================== 잉여 활용 한 장


#: 잉여 시나리오 표의 머리글 (53세션 3-2).
SURPLUS_SCENARIO_HEADER: tuple[str, ...] = ("활용 방안", "연 수익", "비고")

#: 잉여 활용 장의 각주. **자격요건을 판정하지 않는다는 것을 밝힌다.**
#
# **「외부 판매는 단가를 입력해야 산출됩니다」 를 뺐다** (58세션). 단가에
# 기본값이 생겨 그 문장이 늘 참이 아니게 됐고, 비웠을 때는 표의 비고가
# 「판매 단가를 넣으면 산출됩니다」 로 그 자리에서 말한다. 자리에 들어가는
# 것은 :meth:`SurplusResult.applied_price_note` — **무슨 단가로 산출했는가**다.
SURPLUS_PAGE_NOTE = (
    "상계거래는 계약 변경과 역송 계량기가 필요합니다. "
    "제도별 자격요건은 판정하지 않았습니다 — 금액만 참고하십시오."
)

#: 절감액에 이미 든 시나리오에 붙이는 꼬리표 (48세션과 같은 어휘).
SURPLUS_CHOSEN_MARK = "절감액에 반영"


@dataclass(frozen=True)
class SurplusPage:
    """잉여 활용 한 장의 재료 (53세션 3절).

    41세션에 잉여가 **개선안에서 빠져** 태양광 안으로 들어갔다. 그런데 잉여가
    실제로 나면 보여 줄 자리가 없어져, 소형 사무빌딩의 연 23,416 kWh 가 태양광
    장의 **각주 한 줄**로만 나왔다 — 얼마가 언제 남고 그것으로 무엇을 할 수
    있는지가 판단 재료인데도 그렇다.

    **개선안으로 되돌리는 것이 아니다.** 수단 목록(:data:`MEASURE_CATALOG`)은
    여섯 그대로이고, 이 장은 태양광 장 **다음에 붙는 결과 한 장**이다.
    **잉여가 0 이면 만들지 않는다** (:func:`surplus_page` 가 ``None`` 을 낸다).
    """

    lead: str
    facts: tuple[tuple[str, str], ...]
    """지표 넷 — 연간 잉여 · 평일 · 토·일·공휴일 · 잉여 없는 최대 용량."""
    scenario_rows: tuple[tuple[str, ...], ...]
    """시나리오 표. 머리글이 첫 줄이다. **기준선(출력제어)은 싣지 않는다** (27세션).

    언제나 0원이고 「아무것도 하지 않는다」 는 표에 세울 방안이 아니다 — 그것이
    기준선이라는 사실은 화면 라디오가 고른 자리로 말한다 (57세션)."""
    notes: tuple[str, ...] = ()
    """각주 **줄 목록.** 한 줄에 하나씩 「※」 가 선다 (60세션 1절).

    58세션까지는 한 문자열이었다. 자격요건 각주와 적용 단가 각주를 빈칸으로
    이어 넘기는 바람에 **「※」 가 한 줄 가운데 또 섰다** — 뒤 문장이 제 표식을
    달고 오기 때문이다. 성격이 다른 두 말이므로 **줄을 가른다.**"""
    figure: bytes | None = None
    figure_caption: str = ""


def _share(part: float, whole: float) -> str:
    return f"{part / whole * 100:,.1f}%" if whole else "—"


def surplus_page(
    surplus: SurplusResult | None,
    *,
    capacity_kwp: float,
    surplus_free_kwp: float | None = None,
    chosen_scenario: str = "",
    usage: UsageData | None = None,
    surplus_kw: pd.Series | None = None,
) -> SurplusPage | None:
    """잉여 활용 한 장. **잉여가 없으면 ``None`` 이다** (53세션 3-1).

    대형 샘플은 전량 자가소비라 이 장이 생기지 않는다 — 빈 장을 만들어
    「없습니다」 라고 적는 것은 슬라이드를 한 장 늘리는 일일 뿐이다. 그 사실은
    태양광 장의 결론 한 줄이 이미 말한다 (:func:`_solar_conclusion`).
    """
    if surplus is None or surplus.total_kwh <= 0:
        return None
    off_day_kwh = surplus.weekend_kwh + surplus.holiday_kwh
    facts: list[tuple[str, str]] = [
        ("연간 잉여", f"{surplus.total_kwh:,.0f} kWh"),
        (
            "평일 잉여",
            f"{surplus.weekday_kwh:,.0f} kWh ({_share(surplus.weekday_kwh, surplus.total_kwh)})",
        ),
        (
            "토·일·공휴일 잉여",
            f"{off_day_kwh:,.0f} kWh ({_share(off_day_kwh, surplus.total_kwh)})",
        ),
    ]
    if surplus_free_kwp:
        facts.append(("잉여 없는 최대 용량", f"{surplus_free_kwp:,.0f} kWp"))
    # **셋을 다 싣는다** (59세션 4·13절). 53세션까지는 출력제어를 「기준선」 이라
    # 보고 뺐는데(27세션이 화면에서 뺀 것을 따랐다), **57세션에 그것이 기본
    # 선택이 됐다.** 빼 두면 지금 절감액에 든 것이 표에 없어, 「절감액에 반영」
    # 표식이 어느 줄에도 안 붙고 상계거래가 골라진 것처럼 읽힌다.
    #
    # **금액이 큰 순서로 세운다** (59세션 10절). 58세션에 단가 기본값이 생겨
    # 「미산출」 이 사라지면서 순위가 뒤집혔는데(외부 판매가 상계거래보다 크다),
    # 시나리오 정의 순서가 고정이라 그 사실이 표에서 안 보였다. 금액을 못 낸
    # 줄은 맨 뒤다 — 견줄 수 없는 것을 견주는 자리에 세우지 않는다.
    ordered = sorted(
        surplus.scenarios,
        key=lambda item: (item.revenue_won is None, -(item.revenue_won or 0.0)),
    )
    rows: list[tuple[str, ...]] = [SURPLUS_SCENARIO_HEADER]
    for item in ordered:
        rows.append(
            (
                item.name,
                _won(item.revenue_won) if item.is_priced else _UNPRICED,
                _surplus_remark(surplus, item, chosen=item.name == chosen_scenario),
            )
        )
    figure = None
    if usage is not None and surplus_kw is not None:
        figure = _safe_figure(
            lambda: figures.surplus_daily_png(usage, surplus_kw, size=MEASURE_PAIR_FIGURE),
            "잉여 활용 · 일별 잉여",
        )
    return SurplusPage(
        lead=narrative.surplus_page_lead(
            capacity_kwp=capacity_kwp,
            total_kwh=surplus.total_kwh,
            off_day_share=surplus.off_day_share,
        ),
        facts=tuple(facts),
        scenario_rows=tuple(rows),
        # **적용 단가를 표 아래에 적는다** (58세션). 화면 캡션·Excel 참고사항·
        # Word 부록이 쓰는 것과 **같은 문장**이다 — 값은 실제 적용 단가에서 온다.
        notes=tuple(part for part in (SURPLUS_PAGE_NOTE, surplus.applied_price_note) if part),
        figure=figure,
        figure_caption=_SURPLUS_DAILY_CAPTION,
    )


#: 잉여 그래프 캡션. **화면 그림과 같은 것을 쓴다** (15세션 2-6).
_SURPLUS_DAILY_CAPTION = "일별 잉여 — 토요일·일요일을 색으로 갈랐습니다."


def _surplus_remark(surplus: SurplusResult, scenario: SurplusScenario, *, chosen: bool) -> str:
    """시나리오 한 줄의 비고. **고른 것에는 그 사실을 적는다** (48세션).

    상계거래는 금액만으로는 무엇이 일어나는지 알 수 없다 — 당월 사용량에서
    얼마가 차감되는지가 그 방안의 실체다.

    **조각은 :func:`~kwise.report.narrative.note_line` 이 잇는다** (67세션 3절).
    60세션이 각주를 한 자리로 모으면서 여기만 ``" · ".join`` 으로 남았다 —
    지금 조각들로는 결과가 같지만, **마침표로 끝나는 조각이 붙는 순간** 옛
    자리는 「… 않았습니다**. ·** 다음」 으로 겹친다. 규칙이 한 자리에 있으면
    그날 여기도 함께 고쳐진다.
    """
    parts: list[str] = []
    if scenario.name == surplus_module.OFFSET_SCENARIO and surplus.offset is not None:
        parts.append(f"당월 차감 {surplus.offset.deducted_kwh:,.0f} kWh")
        # **잔여가 0 이어도 적는다** (59세션 10절). 표 아래 각주가 「상계거래 SMP
        # 120원/kWh」 라고 적는데 금액은 그 단가와 무관하다 — 잉여가 사용량에 다
        # 잠겨 **SMP 가 곱해질 몫이 없기** 때문이다. 잔여 줄이 없으면 「120원인데
        # 왜 이 금액인가」 를 묻게 된다.
        parts.append(f"기간 말 잔여 {surplus.offset.remaining_kwh:,.0f} kWh")
    elif not scenario.is_priced:
        parts.append("판매 단가를 넣으면 산출됩니다")
    if chosen:
        parts.append(SURPLUS_CHOSEN_MARK)
    return narrative.note_line(*parts)


#: 「하한이 요금적용전력에 걸리지 않아…」 안내의 사실 ID (48세션에 붙은 것).
#: 계약전력 장이 이 안내를 각주로 옮겨 적는다 (59세션 9절).
CONTRACT_FLOOR_NOT_BINDING_FACT = "contract.floor_not_binding"


def measure_entries(
    *,
    switch: TariffSwitchResult | None = None,
    contract: ContractAdjustment | None = None,
    demand_response: DemandResponseResult | None = None,
    power_factor: PowerFactorResult | None = None,
    solar: SolarPoint | None = None,
    solar_certainty: Certainty | None = None,
    solar_notices: tuple[Notice, ...] = (),
    solar_unpriced_reason: str = "",
    ess: EssResult | None = None,
    ess_curve: EssTargetCurve | None = None,
    ess_optimum: EssOptimum | None = None,
    surplus: SurplusResult | None = None,
    usage: UsageData | None = None,
    day: RepresentativeDay | None = None,
    dr_profile: DrProfile | None = None,
    solar_generation_kw: pd.Series | None = None,
    surplus_kw: pd.Series | None = None,
    surplus_free_kwp: float | None = None,
    area_m2: float | None = None,
    base_fee_months: float | None = None,
) -> tuple[MeasureEntry, ...]:
    """검토한 수단을 **7장 순서 그대로** 항목으로 만든다.

    차트 재료(``usage`` 이하)를 주면 수단마다 그림을 함께 굽는다 (15세션 2절).
    주지 않으면 표만 나온다 — 그림이 없다고 보고서가 실패하지 않는다.

    주지 않은 수단은 만들지 않는다 — 켜지 않은 수단이 보고서에 들어가면
    "검토하지 않은 것" 이 "검토했더니 이만큼" 으로 둔갑한다.
    """
    entries: dict[str, MeasureEntry] = {}

    if switch is not None:
        now_option = switch.current.selection.option
        best_option = switch.best.selection.option
        entries["tariff_switch"] = MeasureEntry(
            kind=measure_kind("tariff_switch"),
            conclusion=(
                f"선택{now_option} → 선택{best_option} 로 바꾸면 "
                f"{_won(switch.saving_won)} 줄어듭니다."
                if switch.switch_needed
                else f"현행 선택{now_option} 이 이미 최선입니다. 바꿀 이유가 없습니다."
            ),
            saving=_measure_saving(switch.annual_saving_won, switch.saving_won),
            saving_annual=_annual_saving(switch.annual_saving_won, switch.saving_won),
            has_saving=bool(switch.saving_won),
            investment=_won(0.0),
            payback=_payback_text(0.0, 0.0),
            certainty=str(switch.certainty),
            cautions=(
                "설비 도입과 무관한 확정 계산입니다. 감도를 적용하지 않습니다.",
                *body_lines(switch.notices),
            ),
            notices=switch.notices,
            figure=_safe_figure(
                lambda: figures.tariff_option_png(switch, size=TARIFF_FIGURE),
                "선택요금 전환 · 요금제별 구성",
            ),
            figure_caption="요금제별 요금 구성과 현행 대비 차액",
        )

    if contract is not None:
        entries["contract"] = MeasureEntry(
            kind=measure_kind("contract"),
            conclusion=_contract_conclusion(contract),
            # **0원이 아니라 결론이다** (48세션). 하한이 안 걸려 줄어들 몫 자체가
            # 없는 자리는 「0원」 이 계산이 덜 된 것처럼 읽힌다 — 화면과 같은
            # 말을 쓴다.
            saving=_contract_saving(contract, contract.saving_won),
            saving_annual=_contract_saving(contract, contract.annual_saving_won),
            has_saving=bool(contract.saving_won),
            investment=_won(0.0),
            # **회수기간은 Excel 과 같은 말이다** (83세션 13). 절감이 없는데
            # 「즉시」 라 적으면 즉시 회수된다고 읽힌다.
            payback=_payback_text(0.0, 0.0) if contract.saving_won else "—",
            certainty=str(contract.certainty),
            # **같은 문장을 두 번 싣지 않는다** (102세션 4절). `MARGIN_NOTICE`
            # 가 `CONTRACT_CHANGE_WARNING` 과 **글자까지 같은 사본**이라, 앞에
            # 세운 한 줄과 `contract.notices` 에서 온 한 줄이 7.2 주의사항에
            # **잇달아 두 번** 섰다 (84·100세션). **화면과 같은 방식이다** —
            # `ui\views\measures.py` 가 같은 문자열을 걸러 내고 있다.
            # **사본 둘을 합치는 것은 여기서 안 한다** — 쓰는 자리 여섯을 함께
            # 옮기는 일이라, 뿌리는 미해결에 이름으로 남겼다.
            cautions=(
                CONTRACT_CHANGE_WARNING,
                *(
                    line
                    for line in body_lines(contract.notices)
                    if line != CONTRACT_CHANGE_WARNING
                ),
            ),
            notices=contract.notices,
            # **여지가 없으면 왜 없는지 보인다** (39세션 4-2). 화면이 83세션에
            # 세운 지표와 같은 값이다.
            actionable=contract.reducible,
            # **근거가 결과보다 먼저다** (53세션 6-1).
            facts_first=True,
            # **판정을 가르는 세 수다** (83세션). 화면과 같은 자리·같은 이름이다.
            facts=_contract_facts(contract),
            figure=_safe_figure(
                lambda: figures.contract_headroom_png(contract, size=MEASURE_STRIP_FIGURE),
                "계약전력 조정 · 하한 판정",
            ),
            figure_caption=_CONTRACT_HEADROOM_CAPTION,
            # **「없음」 옆에 까닭이 서야 한다** (59세션 9절). 왜 안 주는지를
            # 계산이 이미 안내로 내고 있었고(화면 판정 줄에 있다), 슬라이드만
            # 그것을 안 읽었다. **문장을 새로 짓지 않는다.**
            slide_note=_notice_text(contract.notices, CONTRACT_FLOOR_NOT_BINDING_FACT),
        )

    if demand_response is not None:
        entries["demand_response"] = MeasureEntry(
            kind=measure_kind("demand_response"),
            conclusion=_dr_conclusion(demand_response, dr_profile),
            saving=(
                # 사유는 이미 「미산출 — …」 꼴이다. 앞에 한 번 더 붙이지 않는다
                # (28세션 1-4 에 문구를 줄이면서 겹침이 드러났다).
                _won(demand_response.settlement_won)
                if demand_response.is_priced
                else demand_response.settlement_label
            ),
            investment=_won(0.0),
            # **Excel 과 같은 말이다** (83세션 13). 절감이 없는데 「즉시」 라
            # 적으면 즉시 회수된다고 읽힌다 — 같은 표의 같은 칸이므로 함께 맞췄다.
            payback=(
                _payback_text(0.0, 0.0)
                if demand_response.is_priced
                else UNPRICED_REASONS["no_saving"]
            ),
            certainty=str(demand_response.certainty),
            has_saving=demand_response.is_priced and bool(demand_response.settlement_won),
            cautions=(
                "정산 단가는 전력거래소 월별 순편익가격과 사업자 수수료로 정해집니다. "
                "**수요관리사업자 상담이 필요합니다.**",
                *body_lines(demand_response.notices),
            ),
            notices=demand_response.notices,
            slide_note=_notice_text(demand_response.notices, DR_OFF_DAYS_FACT),
            figure=(
                _safe_figure(
                    lambda: figures.dr_daily_png(dr_profile, size=MEASURE_FULL_FIGURE),
                    "경제성DR · 일별 평균 부하",
                )
                if dr_profile is not None
                else None
            ),
            figure_caption="일별 운영시간대 평균 부하 — 붉은 선 아래가 감축 가능일입니다.",
        )

    if power_factor is not None:
        # **화면은 그림이 둘이다** (38세션 3-1). 전력삼각형만 실으면 「어느
        # 시간대가 요금 대상인가」 가 슬라이드에서 빠진다.
        triangle = _safe_figure(
            lambda: figures.power_triangle_png(power_factor), "역률 개선 · 전력삼각형"
        )
        pf_day = (
            _safe_figure(
                lambda: figures.power_factor_day_png(
                    usage,
                    day,
                    current_pct=power_factor.current_pct,
                    target_pct=power_factor.target_pct,
                    size=MEASURE_PAIR_FIGURE,
                ),
                "역률 개선 · 대표일 부하",
            )
            if usage is not None and day is not None
            else None
        )
        entries["power_factor"] = MeasureEntry(
            kind=measure_kind("power_factor"),
            conclusion=_power_factor_conclusion(power_factor),
            saving=_measure_saving(power_factor.annual_saving_won, power_factor.saving_won),
            saving_annual=_annual_saving(power_factor.annual_saving_won, power_factor.saving_won),
            has_saving=bool(power_factor.saving_won),
            investment=_won(power_factor.investment_won),
            payback=_payback_text(power_factor.payback_years, power_factor.investment_won),
            certainty=str(power_factor.certainty),
            cautions=body_lines(power_factor.notices),
            notices=power_factor.notices,
            figure=triangle,
            figure_caption=_PF_TRIANGLE_CAPTION,
            figures=_pair(
                (triangle, _PF_TRIANGLE_CAPTION),
                (pf_day, _PF_DAY_CAPTION),
            ),
        )

    if solar is not None:
        # **화면은 연간 발전량과 대표일 곡선을 함께 낸다** (38세션 3-2).
        # 「한 해에 얼마나 만드나」 와 「그날 피크가 얼마나 내려가나」 는 다른
        # 물음이라, 하나만 실으면 나머지 하나를 슬라이드에서 답하지 못한다.
        solar_annual = solar_day = None
        if usage is not None and solar_generation_kw is not None:
            load, generation = usage, solar_generation_kw
            solar_annual = _safe_figure(
                lambda: figures.solar_annual_png(load, generation, size=MEASURE_PAIR_FIGURE),
                "태양광 · 일별 발전량",
            )
            if day is not None:
                point = day
                solar_day = _safe_figure(
                    lambda: figures.solar_day_png(
                        load, generation, point, size=MEASURE_PAIR_FIGURE
                    ),
                    "태양광 · 대표일 부하",
                )
        # **19세션이 다섯 수단에 붙인 ``*body_lines(*.notices)`` 가 여기만 빠져
        # 있었다** (79세션 1절). 그때 다섯은 옮길 ``*warnings`` 가 있었는데
        # 태양광은 9세션에 선 손수 만든 목록이라 손이 안 갔고, 41세션이 잉여
        # 항목을 태양광 안으로 들이면서 그 항목이 갖고 있던
        # ``*body_lines(surplus.notices)`` 마저 사라졌다. **의도가 아니라 이관
        # 누락이다** — 이 목록은 39세션 4-2 보다 앞선다.
        cautions = [
            "발전량 예측은 피크 발전량을 과소 산출하는 경향이 있어 피크 절감량이 "
            "보수적으로 나옵니다.",
        ]
        if solar.investment_won is None and solar_unpriced_reason:
            cautions.append(solar_unpriced_reason)
        # **잉여를 여기로 녹였다** (41세션 2-1·2-6). 7.7 은 개선안이 아니라
        # 태양광의 결과다 — 장을 따로 세우면 같은 발전량 이야기를 두 번 한다.
        #
        # **잉여가 0 이면 아무것도 붙이지 않는다** (39세션 4-2). 팔 것이 없는데
        # 파는 절차를 조심하라는 말이 결론보다 길어진다.
        surplus_facts: tuple[tuple[str, str], ...] = ()
        surplus_note = ""
        if surplus is not None and surplus.total_kwh > 0:
            # **파는 길만 적는다** — 출력제어에는 자격요건이 없고, 상계거래는
            # 1,000 kW 를 넘으면 아예 선택지에 없다 (41세션 2-3).
            sellable = " · ".join(
                item.name
                for item in surplus.scenarios
                if item.name != surplus_module.CURTAIL_SCENARIO
            )
            cautions.append(
                f"{sellable}의 **자격요건은 판정하지 않았습니다.** 금액만 참고하십시오."
            )
            surplus_facts = (
                ("자가소비", f"{(surplus.generation_kwh - surplus.total_kwh) / 1000:,.1f} MWh"),
                ("잉여", f"{surplus.total_kwh / 1000:,.1f} MWh"),
            )
            # **시나리오 줄을 각주에서 뺐다** (53세션 3절). 잉여가 나면
            # **「잉여 활용」 장이 다음에 붙어** 같은 것을 표로 낸다 — 각주가
            # 「상계거래 … · 외부 판매 … · **기준선 0원**」 을 이어 적고 있었고,
            # 그 줄은 27세션에 화면에서 뺀 것이다.
            #
            # **대신 절감액이 무엇의 합인지 적는다** (59세션 5절). 큰 글씨
            # 절감액에 잉여 수익이 얹혀 있는데 그 사실이 어디에도 없었다 —
            # 화면은 절감액 물음표가 늘 이 줄을 낸다 (57세션). 잉여 장과
            # 겹치지 않는다: 그쪽은 **어느 것을 골랐나**를 표식으로 말한다.
            surplus_note = narrative.solar_saving_breakdown(
                self_consumption_won=solar.self_consumption_saving_won,
                surplus_scenario=solar.surplus_scenario,
                surplus_revenue_won=solar.surplus_revenue_won,
            )
        # **역률 영향을 큰 글자에 녹이지 않는다** (59세션 12절 · 목록 P6).
        # 태양광이 유효전력만 상쇄해 역률이 떨어지고 역률요금이 는다 — 사실이고
        # 계산이 이미 내고 있다(``power_factor_extra_won``). 그러나 카드의
        # 절감액은 「그 수단만 적용했을 때」 여야 하므로(31세션) **둘을 나눠
        # 보인다.** Word 는 주의사항 목록이, PPT 는 각주가 받는다 —
        # **같은 문장**이다.
        #
        # **역률을 함께 넘긴다** (79세션 1절). 금액만으로는 「왜 늘었나」 를 못
        # 말하는데, 덱에서 이 각주가 태양광 장이 역률을 말하는 유일한 자리다 —
        # 주의사항 표는 그림 굽기 실패 시의 폴백이라 서지 않는다 (60세션 10절).
        # 화면 2단계와 Excel 은 이미 제 자리에서 내고 있어 값을 주지 않는다.
        power_factor_line = narrative.power_factor_adjusted_saving(
            saving_won=(solar.annual_saving_won if base_fee_months else solar.total_saving_won),
            extra_won=(
                annualize(solar.power_factor_extra_won, base_fee_months)
                if base_fee_months
                else solar.power_factor_extra_won
            ),
            after_pct=solar.power_factor_after_pct,
        )
        if power_factor_line:
            cautions.append(power_factor_line)
            surplus_note = narrative.note_line(surplus_note, power_factor_line)
        entries["solar"] = MeasureEntry(
            kind=measure_kind("solar"),
            conclusion=_solar_conclusion(
                solar,
                surplus_free_kwp,
                area_m2,
                # 잉여가 나면 **「잉여 활용」 장이 뒤따른다** (53세션 3-1).
                surplus_page_follows=surplus is not None and surplus.total_kwh > 0,
            ),
            saving=_measure_saving(solar.annual_saving_won, solar.total_saving_won),
            saving_annual=_annual_saving(solar.annual_saving_won, solar.total_saving_won),
            has_saving=bool(solar.total_saving_won),
            investment=(
                _won(solar.investment_won)
                if solar.investment_won is not None
                else f"{_UNPRICED} — {solar_unpriced_reason or '단가 미입력'}"
            ),
            payback=_payback_text(solar.payback_years, solar.investment_won),
            certainty=str(solar_certainty)
            if solar_certainty is not None
            else str(Certainty.MEDIUM),
            cautions=(*cautions, *body_lines(solar_notices)),
            notices=(*solar_notices, *(surplus.notices if surplus is not None else ())),
            figure=solar_day,
            figure_caption=_SOLAR_DAY_CAPTION,
            figures=_pair(
                (solar_annual, _SOLAR_ANNUAL_CAPTION),
                (solar_day, _SOLAR_DAY_CAPTION),
            ),
            facts=(
                *(
                    (("잉여 없는 최대 용량", f"{surplus_free_kwp:,.0f} kWp"),)
                    if surplus_free_kwp
                    else ()
                ),
                *surplus_facts,
            ),
            slide_note=surplus_note,
        )

    if ess is not None:
        # **목표가 어디서 나왔는지 답하는 그림이 빠져 있었다** (38세션 3-3).
        # 카드가 낸 목표는 이 U곡선의 최소 지점이다.
        # **회수기간 곡선을 뺐다** (46세션). 개략 산정이라 값이 카드와 달라
        # 한 산출물에 두 숫자가 남았다. 자리에 목표별 사양 표가 들어간다 —
        # 전부 카드 기준 참값이고 정밀화가 이미 잰 점이라 추가 계산이 없다.
        ess_table: tuple[tuple[str, ...], ...] = ()
        ess_caption = ""
        if ess_optimum is not None and ess_curve is not None:
            spec = frames.ess_spec_frame(
                ess_optimum, baseline_demand_kw=ess_curve.baseline_demand_kw
            )
            ess_table = frames.ess_spec_rows(spec)
            # **캡션이 표와 어긋나면 안 된다** (59세션 3절). 계약전력 과다 자료는
            # 저감량이 전 줄 0 kW 인데 「목표를 낮추면 저감량은 는다」 고 적혔다.
            ess_caption = frames.ess_spec_caption(spec)
        # **표가 위에 서면 그림은 납작하게 굽는다** (53세션 2절). 남는 높이가
        # 1in 남짓이라 세로가 긴 비율로는 폭을 못 쓴다.
        ess_size = MEASURE_STRIP_FIGURE if ess_table else MEASURE_PAIR_FIGURE
        ess_day = (
            _safe_figure(
                lambda: figures.ess_day_png(usage, ess.dispatch, day, size=ess_size),
                "ESS · 대표일 부하",
            )
            if usage is not None and day is not None
            else None
        )
        entries["ess"] = MeasureEntry(
            kind=measure_kind("ess"),
            conclusion=_ess_conclusion(ess),
            saving=_measure_saving(ess.annual_saving_won, ess.total_saving_won),
            saving_annual=_annual_saving(ess.annual_saving_won, ess.total_saving_won),
            has_saving=bool(ess.total_saving_won),
            investment=_won(ess.investment_won),
            payback=_payback_text(ess.payback_years, ess.investment_won),
            certainty=str(ess.certainty),
            cautions=(
                "규칙기반 단일 디스패치이며 OPEX·열화·교체비를 넣지 않은 단순 회수기간입니다.",
                *body_lines(ess.notices),
            ),
            notices=ess.notices,
            figure=ess_day,
            figure_caption=_ESS_DAY_CAPTION,
            figures=_pair((ess_day, _ESS_DAY_CAPTION)),
            spec_table=ess_table,
            spec_caption=ess_caption,
            spec_note=spec_mark_note(row[-1] for row in ess_table[1:]),
        )
    elif ess_optimum is not None and ess_optimum.below_minimum and ess_curve is not None:
        # **최소 규격에 못 미치면 사양 표를 싣지 않는다** (50세션 3-3). 회수기간도
        # 목표별 사양도 낼 것이 없다 — 살 물건이 없기 때문이다. 대신 이 건물의
        # 사실인 필요 출력·용량·방전시간은 낸다. **확인되지 않은 판정을 적지 않는다.**
        entries["ess"] = MeasureEntry(
            kind=measure_kind("ess"),
            conclusion=body_lines(ess_optimum.notices)[0]
            if ess_optimum.notices
            else NOT_VIABLE_CONCLUSION,
            saving=f"{_UNPRICED} — 최소 규격에 못 미쳐 사양을 정하지 않았습니다.",
            saving_annual=f"{_UNPRICED} — 최소 규격에 못 미쳐 사양을 정하지 않았습니다.",
            has_saving=False,
            investment=f"{_UNPRICED} — 사양 미정",
            payback=_UNPRICED,
            certainty=str(Certainty.HIGH),
            notices=ess_optimum.notices,
            actionable=False,
            facts=(
                ("요금적용전력", f"{ess_curve.baseline_demand_kw:,.0f} kW"),
                ("필요 출력", f"{ess_optimum.required_power_kw:,.1f} kW"),
                ("필요 용량", f"{ess_optimum.required_capacity_kwh:,.1f} kWh"),
                ("방전시간", f"{ess_optimum.required_discharge_hours:,.2f}h"),
                ("상업용 최소 규격", f"{ess_optimum.minimum_power_kw:,.0f} kW"),
            ),
        )
    elif ess_optimum is not None and not ess_optimum.viable and ess_curve is not None:
        # **켠 수단이 산출물에서 사라지면 안 된다** (48세션). 성립하는 목표가
        # 없어 카드가 목표를 제시하지 않은 경우인데, 장을 통째로 빼면 「검토하지
        # 않았다」 와 구분되지 않는다. 계약전력 「하향 여지 없음」·잉여 0 과 같은
        # 틀로 적는다 — 결론과 **왜 없는지 보이는 지표**, 실행 주의사항은 없다.
        #
        # **성립하지 않는 까닭이 둘이다** (59세션 2절 · 56세션).
        #
        #     ① 마진 조건   창을 훑어 잰 점이 있고, 그 점들이 전부 회수되지 않는다
        #     ② 갑 종별     기본요금이 계약전력에 붙어 **훑기 전에 끊는다** — 잰 점이 없다
        #
        # 잰 점이 없으면 표를 그리지 않는다 — 화면 카드와 같은 규약이다
        # (:mod:`kwise.ui.views.measures`). 결론은 계산이 낸 안내를 그대로
        # 옮기므로 갑 종별에서는 :data:`BASE_FEE_ON_CONTRACT_CONCLUSION` 이
        # 실린다 — 문장을 여기서 새로 짓지 않는다.
        measured = bool(ess_optimum.points)
        margin_spec = (
            frames.ess_spec_frame(ess_optimum, baseline_demand_kw=ess_curve.baseline_demand_kw)
            if measured
            else None
        )
        lines = body_lines(ess_optimum.notices)
        reason = "성립하는 목표가 없어" if measured else "성립하지 않아"
        facts = (
            (
                ("요금적용전력", f"{ess_curve.baseline_demand_kw:,.0f} kW"),
                ("성립 한계 방전시간", f"{ess_curve.viable_limit_hours:,.2f}h"),
                ("가장 짧은 방전시간", f"{_shortest_discharge_hours(ess_curve):,.2f}h"),
            )
            if measured
            # 갑 종별은 이 둘이 까닭 전부다 — 기본요금이 붙는 자리가 피크가 아니다.
            else (
                ("요금적용전력", f"{ess_curve.baseline_demand_kw:,.0f} kW"),
                ("기본요금 기준", "계약전력"),
            )
        )
        entries["ess"] = MeasureEntry(
            kind=measure_kind("ess"),
            conclusion=lines[0] if lines else NOT_VIABLE_CONCLUSION,
            saving=f"{_UNPRICED} — {reason} 사양을 정하지 않았습니다.",
            saving_annual=f"{_UNPRICED} — {reason} 사양을 정하지 않았습니다.",
            has_saving=False,
            investment=f"{_UNPRICED} — 사양 미정",
            payback=_UNPRICED,
            certainty=str(Certainty.HIGH),
            notices=ess_optimum.notices,
            actionable=False,
            facts=facts,
            spec_table=frames.ess_spec_rows(margin_spec) if margin_spec is not None else (),
            spec_caption=frames.ess_spec_caption(margin_spec) if margin_spec is not None else "",
        )

    return tuple(entries[kind.key] for kind in MEASURE_CATALOG if kind.key in entries)


# ===================================================================== 보고서 재료


@dataclass(frozen=True)
class DocumentSections:
    """보고서 한 벌의 재료.

    ``measures`` 가 비면 3·4장을 만들지 않는다 — 진단만 보는 정상 경로다.
    """

    usage: UsageData
    bill: BillingResult
    diagnosis: Diagnosis | None = None
    comparison: ComparisonResult | None = None
    sensitivity: tuple[SensitivityRange, ...] = ()
    measures: tuple[MeasureEntry, ...] = ()
    building_name: str = ""
    prepared_on: dt.date | None = None
    reviewed_labels: tuple[str, ...] = ()
    skipped_labels: tuple[str, ...] = ()
    worksheets: tuple[Worksheet, ...] = ()
    """계산 근거 표 (22세션 2절). 화면 카드가 접어 둔 것과 **같은 표**다."""
    tariff_table: TariffTable | None = None
    """부록 B 의 요금표 줄. 없으면 그 줄만 빠진다."""
    ess_cases: pd.DataFrame | None = None
    """ESS 조달 사례. 17세션에 화면에서 뺀 표가 부록 A 로 간다."""
    surplus: SurplusPage | None = None
    """잉여 활용 한 장 (53세션 3절). **잉여가 0 이면 ``None`` 이라 장이 안 생긴다.**"""
    temperature: pd.Series | None = None
    """시간별 기온 (℃). PPT 「전력사용현황 및 부하패턴」이 사용량과 겹쳐 그린다
    (38세션 2-1). **없으면 사용량만 그린다** — 지역은 선택 입력이고, 고른 격자·
    기간을 사전 취득분이 덮지 못할 수도 있다 (화면과 같은 규칙)."""

    @property
    def prepared(self) -> dt.date:
        return self.prepared_on if self.prepared_on is not None else dt.date.today()

    @property
    def building(self) -> str:
        return self.building_name or self.usage.meta.source_name

    def appendix(self) -> tuple[str, ...]:
        """보고서 부록에 실을 **참고 등급** 전부 (19세션 1절).

        화면에서 뺀 문구가 어디에도 남지 않으면 그냥 사라진 것이다. 요금·진단·
        조합·수단이 낸 참고를 한자리에 모아 중복을 지우고 싣는다.
        """
        groups: list[tuple[Notice, ...]] = [self.bill.notices]
        if self.diagnosis is not None:
            groups.append(self.diagnosis.notices)
        if self.comparison is not None:
            groups.append(self.comparison.notices)
        groups.extend(entry.notices for entry in self.measures)
        return tuple(item.text for item in report_appendix(*groups))

    def appendix_data(self) -> AppendixData:
        """부록 A·B·C 재료를 **한 번에** 만든다 (22세션 3절).

        Word 와 Excel 이 같은 것을 실어야 하므로 만드는 자리를 하나로 둔다.
        """
        groups: list[tuple[Notice, ...]] = [self.bill.notices]
        if self.diagnosis is not None:
            groups.append(self.diagnosis.notices)
        if self.comparison is not None:
            groups.append(self.comparison.notices)
        groups.extend(entry.notices for entry in self.measures)
        grounds = tuple(
            (entry.kind.key, texts(report_body(entry.notices), Severity.BASIS))
            for entry in self.measures
            if texts(report_body(entry.notices), Severity.BASIS)
        )
        return AppendixData(
            worksheets=self.worksheets,
            grounds=grounds,
            cases=self.ess_cases,
            limits=known_limits(*groups),
            assumptions_rows=reference_rows(self.tariff_table),
        )

    def scope(self) -> tuple[tuple[str, ...], tuple[str, ...]]:
        """(검토함, 미검토). 넘겨받지 않았으면 수단 목록에서 만든다."""
        if self.reviewed_labels or self.skipped_labels:
            return self.reviewed_labels, self.skipped_labels
        done = {entry.kind.key for entry in self.measures}
        return (
            tuple(kind.title for kind in MEASURE_CATALOG if kind.key in done),
            tuple(kind.title for kind in MEASURE_CATALOG if kind.key not in done),
        )


# ===================================================================== 문서 골격


def _set_korean_font(document: DocumentType) -> None:
    """본문 폰트를 한글 폰트로. **동아시아 글꼴을 따로 지정해야** Word 가 쓴다."""
    # **없을 때의 갈래를 걷어냈다** (60세션 12절). 다섯은 기본 템플릿에 늘 있고,
    # 없으면 `python-docx` 가 `KeyError: "no style with name 'Heading 1'"` 로
    # **어느 스타일인지 짚어 준다** — 조용히 건너뛰면 글꼴이 빠진 문서가 그냥 나간다.
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = KOREAN_FONT
        rpr = style.element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)


def _para(document: DocumentType, text: str = "", *, style: str | None = None) -> Paragraph:
    """문단 하나. **마크다운을 벗겨 적는다** (39세션 2-6).

    Word 도 PPT 와 같이 마크다운을 해석하지 않아 ``**`` 가 글자로 찍힌다. 벗기는
    자리는 :func:`kwise.report.notices.plain_text` 하나이고, 글자를 쓰는 함수가
    그것을 지나는지 시험이 훑는다.
    """
    if style is not None:
        return document.add_paragraph(plain_text(text), style=style)
    return document.add_paragraph(plain_text(text))


def _heading(document: DocumentType, text: str, *, level: int) -> Paragraph:
    """제목 하나. 문단과 같은 이유로 벗겨 적는다."""
    return document.add_heading(plain_text(text), level=level)


def _add_table(
    document: DocumentType, rows: Sequence[Sequence[str]], *, header: bool = True
) -> None:
    """**Word 표 객체**로 넣는다. 이미지로 넣으면 제안서에 복사해 쓸 수 없다 (10.5)."""
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = TABLE_STYLE
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = plain_text(str(value))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if header and row_index == 0:
                        run.font.bold = True
    document.add_paragraph()


def _add_figure(document: DocumentType, png: bytes, caption: str) -> None:
    document.add_picture(io.BytesIO(png), width=_FIGURE_WIDTH)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = document.add_paragraph(plain_text(caption))
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in label.runs:
        run.font.size = Pt(9)
        run.font.italic = True


def _add_bullets(document: DocumentType, lines: Iterable[str]) -> None:
    for line in lines:
        _para(document, line, style="List Bullet")


def _conclusion(document: DocumentType, text: str) -> None:
    """**결론 한 줄.** 절의 첫 문단이며 굵게 쓴다 — 데이터를 앞에 놓지 않는다."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(plain_text(text))
    run.bold = True


# ===================================================================== 장별


def _cover(document: DocumentType, sections: DocumentSections) -> None:
    title = _heading(document, DOCUMENT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = _para(document, sections.building)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    bill = sections.bill
    _add_table(
        document,
        [
            ["항목", "내용"],
            ["건물명", sections.building],
            [
                "분석 기간",
                f"{bill.period_start:%Y-%m-%d} ~ {bill.period_end:%Y-%m-%d} "
                f"({bill.period_days:.0f}일, 기본요금 {bill.base_fee_months:.2f}개월분)",
            ],
            ["작성일", f"{sections.prepared:%Y-%m-%d}"],
            ["적용 요금표 시행일", f"{bill.effective_date}"],
            ["계약종별", f"{bill.contract_label} {bill.voltage_label} 선택{bill.selection.option}"],
        ],
    )
    document.add_page_break()


def _chapter_summary(document: DocumentType, sections: DocumentSections, number: int) -> None:
    """**이 장만 읽어도 판단이 되어야 한다.**"""
    _heading(document, f"{number}장 {CHAPTER_SUMMARY}", level=1)
    diagnosis = sections.diagnosis
    summary = diagnosis.summary if diagnosis is not None else None

    free = summary.no_investment_saving_won if summary is not None else None
    _conclusion(
        document,
        (
            f"투자 없이 {_won(free)} 를 줄일 수 있습니다."
            if free is not None
            else "투자 없이 가능한 절감액은 계약 정보가 있어야 산출됩니다."
        ),
    )

    rows = [["항목", "값"]]
    rows.append(["투자 없이 가능한 절감액", _won(free)])
    if summary is not None:
        rows.append(
            [
                "선택요금 전환",
                _won(summary.tariff_switch_saving_won),
            ]
        )
        rows.append(["계약전력 조정", _won(summary.contract_saving_won)])
        rows.append(["태양광 피크 기여 가능성", str(summary.pv_potential)])

    best = sections.comparison.best if sections.comparison is not None else None
    if best is not None:
        rows.append(["권장 조합", best.name])
        rows.append(["총 절감액", _won(best.saving_won)])
        rows.append(["투자비", _won(best.investment_won)])
        rows.append(["회수기간", _payback_text(best.payback_years, best.investment_won)])
    _add_table(document, rows)

    # **확실성 등급 줄을 뺐다** (53세션 1-4). 화면에서 28세션에 걷어낸 것이
    # 산출물에만 남아 있었다 — 무엇에 대한 등급인지 이름에 없어 읽는 사람이
    # 되물을 수밖에 없는 값이다.
    _para(document, NOT_INCLUDED_NOTICE + " 자세한 한계는 마지막 장에 있습니다.")
    # 금액은 천 원 단위로 절사해 적는다 (14세션 1절). 항목 합과 합계가 어긋날 수 있다.
    _para(document, TRUNCATION_FOOTNOTE)
    document.add_page_break()


def _chapter_diagnosis(document: DocumentType, sections: DocumentSections, number: int) -> None:
    _heading(document, f"{number}장 {CHAPTER_DIAGNOSIS}", level=1)
    diagnosis = sections.diagnosis
    meta = sections.usage.meta

    # ---- 데이터 개요와 품질
    _heading(document, f"{number}.1 데이터 개요와 품질", level=2)
    quality = diagnosis.quality if diagnosis is not None else None
    missing = f"{quality.missing_ratio:.1%}" if quality is not None else "—"
    _conclusion(
        document,
        f"{meta.start:%Y-%m-%d} ~ {meta.end:%Y-%m-%d} 의 {meta.interval_minutes}분 실측 "
        f"{meta.expected_rows:,}구간을 썼습니다. 결측률 {missing} 입니다.",
    )
    rows = [
        ["항목", "값"],
        ["기간", f"{meta.start:%Y-%m-%d} ~ {meta.end:%Y-%m-%d} ({meta.period_days:.0f}일)"],
        ["검침 간격", f"{meta.interval_minutes}분"],
        ["총 사용량", f"{meta.total_kwh / 1000:,.1f} MWh"],
        ["결측", f"{meta.missing_rows:,}구간 ({meta.missing_ratio:.1%}) — 보간하지 않았습니다"],
    ]
    if quality is not None:
        rows.append(["정전 추정", f"{len(quality.outages)}건"])
        if quality.longest_gap is not None:
            gap = quality.longest_gap
            rows.append(
                [
                    "최장 연속 결측",
                    f"{gap.slots:,}구간 ({gap.start:%Y-%m-%d %H:%M} ~ {gap.end:%Y-%m-%d %H:%M})",
                ]
            )
    _add_table(document, rows)

    if diagnosis is None:
        return

    # ---- 부하 패턴
    _heading(document, f"{number}.2 부하 패턴", level=2)
    pattern = diagnosis.pattern
    _conclusion(
        document,
        f"부하율 {(pattern.load_factor or 0) * 100:,.1f}%, "
        f"기저부하 비율 {(pattern.base_load_ratio or 0) * 100:,.1f}% 입니다.",
    )
    _add_table(
        document,
        [
            ["지표", "값", "의미"],
            ["부하율", f"{(pattern.load_factor or 0) * 100:,.1f}%", "평균 ÷ 최대"],
            [
                "기저부하 비율",
                f"{(pattern.base_load_ratio or 0) * 100:,.1f}%",
                "야간 평균 ÷ 주간 평균",
            ],
            [
                "주말 부하 비율",
                f"{(pattern.weekend_ratio or 0) * 100:,.1f}%",
                "주말 평균 ÷ 평일 평균",
            ],
            [
                "운영시간 외 부하 비중",
                f"{(pattern.off_hours_energy_share or 0) * 100:,.1f}%",
                f"운영 {pattern.operating_hours[0]}~{pattern.operating_hours[1]}시 밖",
            ],
        ],
    )
    _add_figure(
        document,
        figures.hourly_profile_png(diagnosis.peak),
        f"그림 {number}-1. 시간대별 평균 부하 프로파일",
    )

    # ---- 피크 특성
    _heading(document, f"{number}.3 피크 특성", level=2)
    peak = diagnosis.peak
    _conclusion(
        document,
        f"관측 최대수요는 {peak.peak_kw:,.1f} kW, 요금적용전력은 "
        f"{peak.billing_demand_kw:,.1f} kW 입니다."
        + (
            " 경부하 시간대의 피크는 요금적용전력이 되지 않습니다."
            if peak.billing_demand_kw < peak.peak_kw * 0.99
            else ""
        ),
    )
    _add_figure(
        document,
        figures.monthly_peak_png(peak),
        f"그림 {number}-2. 월별 최대수요와 요금적용전력",
    )
    _add_figure(
        document,
        figures.top_hour_png(peak),
        f"그림 {number}-3. 상위 {peak.top_n}구간 시각 분포 — 태양광 기여 가능성의 지표",
    )

    # ---- 현재 요금 구조
    structure = diagnosis.structure
    if structure is not None:
        _heading(document, f"{number}.4 현재 요금 구조", level=2)
        # **화면·PPT 와 같은 몫을 센다** (S124 · ②-40). 앞서는 ``base_won``(역률요금
        # 을 뺀 값)을 ``total_won``(담은 값)으로 나누고 있어 **짝이 안 맞았다** —
        # 역률 85% 를 걸면 이 표의 기본 + 전력량이 합계보다 317,220원 모자라고
        # 비중 합이 99.8% 가 된다. 역률요금이 0원인 자료(간주 92%)에서는 값이 그대로다.
        base_won = structure.base_with_power_factor_won
        base_share = structure.base_with_power_factor_share
        _conclusion(
            document,
            f"기본요금이 총액의 {base_share:.1%} 입니다 "
            f"({_won(base_won)} / {_won(structure.total_won)}).",
        )
        share = structure.band_share
        # **초과사용부가금은 붙은 자료에서만 한 칸을 더 쓴다** (S127 2절 · ②-32).
        # 화면과 PPT 는 109세션부터 이 칸을 세우는데 **이 표만 안 세우고 있었다** —
        # 분모 ``total_won`` 은 부가금을 담고 있으므로 안 세우면 **기본 + 전력량이
        # 합계에 못 미친다.** S124 가 고친 역률요금 짝 안 맞음과 같은 모양이다.
        # 부가금이 서는 벌이 저장소에 하나도 없어 **한 번도 그려진 적이 없었다** —
        # S127 이 `large-b-short` 를 지어 그 자리에서 드러났다 (122,451,200원 ·
        # 총액의 3.5%). 0원이면 칸이 안 생기므로 지금까지의 산출물은 그대로다.
        excess_won = structure.excess_won
        _add_table(
            document,
            [
                ["구분", "금액·비중"],
                ["기본요금", f"{_won(base_won)} ({base_share:.1%})"],
                [
                    "전력량요금",
                    f"{_won(structure.energy_won)} ({structure.energy_share:.1%})",
                ],
                *(
                    [
                        [
                            "초과사용부가금",
                            f"{_won(excess_won)} ({excess_won / structure.total_won:.1%})",
                        ]
                    ]
                    if excess_won
                    else []
                ),
                ["경부하 사용량 비중", f"{float(share.get('light', 0.0)):.1%}"],
                ["중간부하 사용량 비중", f"{float(share.get('mid', 0.0)):.1%}"],
                ["최대부하 사용량 비중", f"{float(share.get('peak', 0.0)):.1%}"],
            ],
        )

    # ---- 계약전력 적정성
    adequacy = diagnosis.contract
    if adequacy is not None:
        _heading(document, f"{number}.5 계약전력 적정성", level=2)
        floor_ratio = adequacy.contract_floor_ratio
        # **7.2 와 같은 문장을 쓴다** (100세션). 여기에 따로 적혀 있던 두 갈래는
        # 종별 전환을 몰라 「299 kW 로 낮춰라」 하는 판에서 「적정합니다」 라고
        # 적었다. 어휘를 한 곳에 두면 두 장이 갈릴 수 없다.
        _conclusion(document, _contract_conclusion(adequacy.adjustment))
        _add_table(
            document,
            [
                ["항목", "값"],
                ["계약전력", f"{adequacy.contract_kw:,.0f} kW"],
                [
                    f"계약전력의 {floor_ratio:.0%}" if floor_ratio is not None else "하한",
                    f"{adequacy.floor_kw:,.1f} kW" if adequacy.floor_kw is not None else "—",
                ],
                ["최대수요", f"{adequacy.billing_demand_kw:,.1f} kW"],
                ["이용률", f"{adequacy.utilization:.1%}"],
                [
                    "목표 계약전력",
                    f"{adequacy.target_contract_kw:,.0f} kW"
                    if adequacy.target_contract_kw is not None
                    else NO_SAVING,
                ],
                [
                    "예상 절감액",
                    _contract_adequacy_saving(adequacy),
                ],
            ],
        )
        if adequacy.reducible:
            _para(document, CONTRACT_CHANGE_WARNING)
    document.add_page_break()


def _chapter_measures(document: DocumentType, sections: DocumentSections, number: int) -> None:
    _heading(document, f"{number}장 {CHAPTER_MEASURES}", level=1)
    _para(
        document,
        "검토한 수단만 싣습니다. 투자비 순이며, 보지 않은 수단은 마지막 장의 "
        "「검토 범위」에 있습니다.",
    )
    for index, entry in enumerate(sections.measures, start=1):
        _heading(document, f"{number}.{index} {entry.title}", level=2)
        _conclusion(document, entry.conclusion)
        _add_table(
            document,
            [
                ["항목", "값"],
                ["절감액", entry.saving],
                ["투자비", entry.investment],
                ["회수기간", entry.payback],
            ],
        )
        # **표가 있으면 그림보다 먼저다** (46세션). ESS 의 목표별 사양 표가
        # 「이 목표는 어디서 나왔나」 에 답하므로 그림보다 앞에 와야 한다.
        if entry.spec_table:
            _add_table(document, [list(row) for row in entry.spec_table])
            if entry.spec_caption:
                _para(document, entry.spec_caption)
        if entry.figure is not None:
            _add_figure(document, entry.figure, f"그림 {number}-{index}. {entry.figure_caption}")
        if entry.cautions:
            _para(document, "주의사항")
            _add_bullets(document, entry.cautions)
    document.add_page_break()


def _chapter_comparison(document: DocumentType, sections: DocumentSections, number: int) -> None:
    _heading(document, f"{number}장 {CHAPTER_COMPARISON}", level=1)
    comparison = sections.comparison
    assert comparison is not None  # build_document 가 없으면 이 장을 부르지 않는다
    best = comparison.best
    # **조합에 무엇이 들었는지 이름 대신 구성으로 적는다** (S112 5절 · ⑱).
    # 조합 이름은 「+ ESS 목표 5,170 kW」 처럼 직전 조합에 더한 것만 말해
    # **어느 요금제로 낸 값인지가 안 보인다** — 조합이 조합 부하에서 선택요금을
    # 다시 고르게 되면서 그 값이 2단계 권고와 다를 수 있다. PPT 15장이 이미
    # 같은 것을 적고 있다. **문구를 늘리지 않았다** — 같은 문장의 한 조각이다.
    baseline = comparison.combinations[0].selection if comparison.combinations else None
    _conclusion(
        document,
        f"권장안은 「{best.composition(baseline)}」 입니다. {_won(best.saving_won)} 를 줄이고 "
        f"투자비는 {_won(best.investment_won)}, 회수기간은 "
        f"{_payback_text(best.payback_years, best.investment_won)} 입니다.",
    )
    rows = [["조합", "요금제", "절감액", "투자비", "회수기간"]]
    for item in comparison.combinations:
        rows.append(
            [
                item.name,
                option_label(item.selection.option),
                _won(item.saving_won),
                _won(item.investment_won),
                _payback_text(item.payback_years, item.investment_won),
            ]
        )
    _add_table(document, rows)
    _para(
        document,
        "**조합마다 요금을 다시 계산했습니다.** 수단별 절감액의 단순 합이 아닙니다 — "
        "태양광이 사용량을 줄이면 최적 선택요금이 바뀌고, ESS 가 피크를 낮추면 "
        "기본요금 기반이 달라집니다.",
    )
    _add_figure(
        document, figures.combination_png(comparison), f"그림 {number}-1. 조합별 절감액과 투자비"
    )

    if sections.sensitivity:
        _heading(document, f"{number}.1 감도", level=2)
        _conclusion(
            document,
            "태양광 출력의 첨예도만 흔들어 본 범위입니다. 요금제 전환·계약전력 조정·"
            "역률 개선은 확정 계산이라 감도를 적용하지 않습니다.",
        )
        # **범위로 적는다. 3열 나열을 하지 않는다** (9.2).
        rows = [["지표", "기준값과 범위"]]
        rows.extend(
            [item.metric, item.text()] for item in sections.sensitivity if item.base is not None
        )
        _add_table(document, rows)
        _para(document, SCENARIO_NAME_CAVEAT)
    document.add_page_break()


def _chapter_scope(document: DocumentType, sections: DocumentSections, number: int) -> None:
    _heading(document, f"{number}장 {CHAPTER_SCOPE}", level=1)
    reviewed, skipped = sections.scope()
    _conclusion(
        document,
        f"검토한 수단은 {len(reviewed)}개, 보지 않은 수단은 {len(skipped)}개입니다. "
        "**보지 않은 것은 '효과가 없다' 가 아닙니다.**",
    )
    _add_table(
        document,
        [
            ["구분", "수단"],
            ["검토함", ", ".join(reviewed) or "없음"],
            ["미검토", ", ".join(skipped) or "없음"],
        ],
    )

    _heading(document, f"{number}.1 미포함 요금요소", level=2)
    _para(document, NOT_INCLUDED_NOTICE)
    _para(document, TRUNCATION_FOOTNOTE)

    _heading(document, f"{number}.2 계약전력 변경 시 주의", level=2)
    _para(document, CONTRACT_CHANGE_WARNING)

    _heading(document, f"{number}.3 추적성", level=2)
    _add_bullets(document, sections.bill.traceability())
    _add_bullets(document, DATA_SOURCES)

    # **참고 등급은 부록 C 로 옮겼다** (22세션 3절). 19세션에 여기 5.5절을 두었는데
    # 부록 C(알려진 한계)와 성격이 같아 같은 말이 두 곳에 실릴 자리였다.


# ===================================================================== 조립


def _appendix_a(document: DocumentType, sections: DocumentSections) -> None:
    """부록 A — **계산 근거 표와 근거 등급 문구, 그리고 조달 사례.**

    화면은 접힘 안에 두고 보고서는 펼친다. 나중에 혼자 읽는 문서이고, 그때
    묻는 것이 바로 「이 숫자가 어떻게 나왔나」다 (19세션 1절).
    """
    data = sections.appendix_data()
    if not data.worksheets and not data.grounds:
        return
    document.add_page_break()
    _heading(document, APPENDIX_TITLES["A"], level=1)
    _para(document, "화면에서 접어 둔 계산 근거입니다. 산식과 대입한 값을 나란히 실었습니다.")
    grounds = dict(data.grounds)
    for sheet in data.worksheets:
        _heading(document, sheet.title, level=2)
        _add_table(document, [list(COLUMNS), *[list(row) for row in sheet.frame().to_numpy()]])
        lines = grounds.get(sheet.key, ())
        if lines:
            _add_bullets(document, lines)
    if data.cases is not None and not data.cases.empty:
        _heading(document, "ESS 조달 사례", level=2)
        _para(document, "투자비 회귀의 원자료입니다. 화면에서는 뺐고(17세션) 여기에 싣습니다.")
        _add_table(
            document,
            [
                [str(name) for name in data.cases.columns],
                *[[str(value) for value in row] for row in data.cases.to_numpy()],
            ],
        )


def _appendix_b(document: DocumentType, sections: DocumentSections) -> None:
    """부록 B — **기준 데이터에서 만든다.** 손으로 옮겨 적지 않는다."""
    rows = sections.appendix_data().assumptions_rows
    if not rows:
        return
    document.add_page_break()
    _heading(document, APPENDIX_TITLES["B"], level=1)
    _para(
        document,
        "이 산출에 쓴 기준 값입니다. 법령 유래와 우리 판단값을 구분해 실었으며, "
        "값은 기준 데이터 파일에서 그대로 가져옵니다.",
    )
    _add_table(
        document,
        [["구분", "항목", "값", "근거", "확인일"], *[list(row) for row in rows]],
    )


def _appendix_c(document: DocumentType, sections: DocumentSections) -> None:
    """부록 C — 알려진 한계와 전제. **19세션 5.5절이 여기로 왔다.**"""
    lines = sections.appendix_data().limits
    if not lines:
        return
    document.add_page_break()
    _heading(document, APPENDIX_TITLES["C"], level=1)
    _para(
        document,
        "결과를 읽는 데 필요한 한계와 전제입니다. 화면에서는 본문을 가리지 않도록 "
        "빼고 여기에 모았습니다.",
    )
    _add_bullets(document, lines)


def build_document(sections: DocumentSections) -> DocumentType:
    """보고서를 만든다.

    **내용이 있는 장만 넣고 번호를 당긴다.** 수단을 하나도 켜지 않으면 3·4장이
    빠져 요약·진단·검토 범위 셋으로 끝난다 — 진단만 보고 받아 가는 정상 경로다.
    비어 있는 장을 제목만 남겨 두면 "검토했는데 결과가 없다" 로 읽힌다.
    """
    document = Document()
    _set_korean_font(document)
    _cover(document, sections)

    number = 1
    _chapter_summary(document, sections, number)
    number += 1
    _chapter_diagnosis(document, sections, number)
    if sections.measures:
        number += 1
        _chapter_measures(document, sections, number)
    if sections.comparison is not None:
        number += 1
        _chapter_comparison(document, sections, number)
    number += 1
    _chapter_scope(document, sections, number)

    # **부록 셋** (22세션 3절). 본문에서 뺀 것이 여기 모인다.
    _appendix_a(document, sections)
    _appendix_b(document, sections)
    _appendix_c(document, sections)
    return document


def document_path(
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    *,
    prefix: str = "kwise_report",
    now: dt.datetime | None = None,
) -> Path:
    """**날짜·시각 접미사를 붙인다.** Word 가 파일을 열고 있으면 덮어쓰기가 실패한다."""
    stamp = (now if now is not None else dt.datetime.now()).strftime("%Y%m%d_%H%M")
    return output_dir / f"{prefix}_{stamp}.docx"


def export_document(
    sections: DocumentSections,
    *,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    prefix: str = "kwise_report",
    now: dt.datetime | None = None,
) -> Path:
    """파일로 쓴다."""
    path = document_path(output_dir, prefix=prefix, now=now)
    path.parent.mkdir(parents=True, exist_ok=True)
    build_document(sections).save(str(path))
    return path


def document_bytes(
    sections: DocumentSections, *, now: dt.datetime | None = None, prefix: str = "kwise_report"
) -> tuple[bytes, str]:
    """내려받기용 바이트와 파일명. **디스크에 남기지 않는다** (10.2)."""
    buffer = io.BytesIO()
    build_document(sections).save(buffer)
    return buffer.getvalue(), document_path(Path(), prefix=prefix, now=now).name
